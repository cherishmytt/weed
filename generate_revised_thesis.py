#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成修改后的激光除草运行监测系统毕业论文
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_revised_thesis():
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.5
    doc.styles['Normal'].paragraph_format.first_line_indent = Pt(24)

    # 辅助函数
    def set_font(run, name, size, bold=False):
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        run.font.size = size
        run.font.bold = bold

    def add_heading1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '黑体', Pt(16), True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Pt(0)
        return p

    def add_heading2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '黑体', Pt(14), True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(0)
        return p

    def add_heading3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '黑体', Pt(12), True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(0)
        return p

    def add_text(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def add_caption(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '宋体', Pt(10.5))
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Pt(0)
        return p

    def add_ref(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '宋体', Pt(10.5))
        p.paragraph_format.first_line_indent = Pt(-21)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.space_after = Pt(3)
        return p

    def add_mermaid_diagram(title, code):
        """添加Mermaid图描述"""
        p = doc.add_paragraph()
        run = p.add_run(f"【{title} - Mermaid语法描述】")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 139)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.first_line_indent = Pt(0)

        code_p = doc.add_paragraph()
        code_run = code_p.add_run(code)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)
        code_p.paragraph_format.left_indent = Pt(42)
        code_p.paragraph_format.line_spacing = 1.2
        code_p.paragraph_format.first_line_indent = Pt(0)

    # ==================== 论文标题 ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('激光除草机器人运行监测系统')
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_p.paragraph_format.space_after = Pt(24)
    title_p.paragraph_format.first_line_indent = Pt(0)

    # ==================== 摘要 ====================
    add_heading2('摘要')

    add_text('杂草与农作物竞争光照、水分及土壤养分，是制约作物产量与品质的重要因素之一。传统化学除草方式虽然效率较高，但长期使用易导致土壤退化、水体污染以及农药残留等环境与食品安全问题。随着计算机视觉与精准农业技术的发展，基于激光的物理除草方案因其非接触、无污染、可精准定位等特性，逐渐成为该领域的研究热点。然而，激光除草机器人在实际作业过程中，如何对机器人的运行状态、检测结果、激光执行等关键信息进行有效监控与管理，是保障系统稳定运行的重要问题。')

    add_text('本课题针对激光除草机器人作业过程中的监控管理需求，设计并实现了一套基于Web的激光除草机器人运行监测系统。系统采用B/S架构，通过软件模拟的方式实现机载端与服务端的通信交互，完成机器人运行状态上报、杂草检测结果展示、激光控制指令下发等核心功能。前端采用Vue 3与Element Plus组件库实现数据可视化界面，后端采用Spring Boot框架提供RESTful API服务，通过WebSocket实现实时数据推送。本系统仅负责YOLO模型的集成调用，不涉及硬件相关的开发工作，通过软件仿真的方式验证监控系统的业务流程与技术可行性，为后续硬件系统的集成与部署提供了软件基础。')

    add_text('本课题所设计并实现的激光除草机器人运行监测系统，实现了机器人状态实时监控、检测记录管理、激光控制指令管理等核心功能。系统界面简洁直观，操作便捷，能够有效帮助操作人员远程监控机器人的作业状态，提升激光除草机器人的作业管理效率。')

    kw_p = doc.add_paragraph()
    kw_run = kw_p.add_run('关键词：')
    kw_run.bold = True
    kw_p.add_run('运行监测系统，除草机器人，Web系统，Spring Boot，Vue.js')
    kw_p.paragraph_format.space_before = Pt(12)
    kw_p.paragraph_format.first_line_indent = Pt(0)

    doc.add_page_break()

    # ==================== 英文摘要 ====================
    add_heading2('Abstract')

    add_text('Weeds compete with crops for sunlight, water, and soil nutrients, which is one of the important factors restricting crop yield and quality. Although traditional chemical weeding methods are efficient, long-term use can easily lead to environmental and food safety problems such as soil degradation, water pollution, and pesticide residues. With the development of computer vision and precision agriculture technology, laser-based physical weeding solutions have gradually become a research hotspot in this field due to their non-contact, pollution-free, and precise positioning characteristics. However, during the actual operation of laser weeding robots, how to effectively monitor and manage key information such as robot operation status, detection results, and laser execution is an important issue to ensure stable system operation.')

    add_text('This thesis designs and implements a Web-based laser weeding robot operation monitoring system to address the monitoring and management requirements during the operation of laser weeding robots. The system adopts a B/S architecture, realizes communication interaction between the on-board end and the server through software simulation, and completes core functions such as robot operation status reporting, weed detection result display, and laser control command distribution. The front-end uses Vue 3 and Element Plus component library to implement a data visualization interface, and the back-end uses Spring Boot framework to provide RESTful API services, realizing real-time data push through WebSocket. This system is only responsible for the integration and invocation of YOLO model, and does not involve hardware-related development work. Through software simulation, the business process and technical feasibility of the monitoring system are verified, providing a software foundation for subsequent hardware system integration and deployment.')

    add_text('The laser weeding robot operation monitoring system designed and implemented in this thesis realizes core functions such as real-time robot status monitoring, detection record management, and laser control command management. The system interface is simple, intuitive, and easy to operate, which can effectively help operators remotely monitor the robot\'s operation status and improve the operation management efficiency of laser weeding robots.')

    kw_p = doc.add_paragraph()
    kw_run = kw_p.add_run('Key Words: ')
    kw_run.bold = True
    kw_p.add_run('Operation monitoring system, Weeding robot, Web system, Spring Boot, Vue.js')
    kw_p.paragraph_format.space_before = Pt(12)
    kw_p.paragraph_format.first_line_indent = Pt(0)

    doc.add_page_break()

    # ==================== 第1章 绪论 ====================
    add_heading1('第1章 绪论')

    add_heading2('1.1 研究背景与意义')

    add_text('农业是国民经济的基础，农作物的产量与品质直接关系到粮食安全与人民生活水平。在农业生产过程中，杂草防治是田间管理的重要环节。据统计，全球每年因杂草造成的农作物产量损失约占总产量的34%，这一比例在发展中国家可能更高[1]。传统的杂草防治方式主要包括人工除草、机械除草和化学除草。人工除草效率低、劳动强度大，无法适应规模化农业生产需求；机械除草虽然效率较高，但容易造成土壤结构破坏和作物损伤；化学除草剂的大量使用会造成土壤污染、水体富营养化、农药残留等生态环境问题，严重威胁农产品质量安全和农业可持续发展[2]。')

    add_text('随着计算机视觉、人工智能与机器人技术的快速发展，精准农业技术为杂草防治提供了新的解决方案。基于机器视觉的杂草检测与物理除草技术，可以实现杂草的精准识别与定点清除，在保证除草效果的同时最大限度减少对环境的影响[3]。其中，激光除草技术因其非接触、无污染、作用时间短、可精准控制等优点，成为近年来精准农业领域的研究热点。激光除草的基本原理是利用高能量激光束照射杂草的生长点，通过热效应使杂草细胞组织坏死，从而达到除草目的[4]。')

    add_text('激光除草机器人是一种集成了机器视觉、运动控制、激光技术等多学科技术的智能农业装备。机器人在实际作业过程中需要对运行状态、检测结果、激光执行情况等进行实时监控与管理。传统的监控方式依赖于本地控制台显示，无法实现远程监控与集中管理，操作人员必须在田间现场才能了解机器人的运行状况[5]。')

    add_text('Web技术的发展为远程监控系统提供了成熟的技术方案。基于B/S架构的Web监控系统，用户只需通过浏览器即可访问系统，无需安装专用客户端，具有跨平台、易部署、维护方便等优点[6]。WebSocket等实时通信技术的应用，使得Web系统能够实现数据的实时推送，满足监控系统的实时性需求。')

    add_text('本课题针对激光除草机器人的监控管理需求，设计并实现一套基于Web的激光除草机器人运行监测系统。系统采用软件仿真的方式，在不依赖实际硬件设备的情况下，验证激光除草机器人监控系统的业务流程与技术可行性。本研究对于推动激光除草技术的实用化进程具有重要的理论意义和应用价值。')

    add_heading2('1.2 国内外研究现状')

    add_text('激光除草技术自20世纪80年代开始研究，进入21世纪后，随着计算机视觉与机器人技术的发展，逐渐向自动化、智能化方向发展[7]。国外方面，美国Carbon Robotics公司开发的Autonomous LaserWeeder机器人已实现商业化应用，该机器人配备了8台高功率激光器和计算机视觉系统，每小时可处理20万株杂草，除草精度达到99%以上，目前已在美国多个大型农场投入使用[8]。荷兰蓝洞公司（Blue River Technology）开发的See & Spray系统采用计算机视觉技术识别杂草，通过精准喷洒除草剂的方式减少农药使用量，该系统已被John Deere公司收购并实现商业化[9]。')

    add_text('国内方面，中国农业大学、浙江大学、南京农业大学等高校开展了激光除草相关研究。中国农业大学现代农业装备实验室在激光除草机器人的视觉检测方面开展了深入研究，提出了基于YOLO的杂草实时检测算法，在田间复杂背景下取得了较好的检测效果[10]。浙江大学农业工程研究所研究了激光参数对不同杂草的杀伤效果，建立了激光除草效果评估模型，为激光除草系统的参数优化提供了理论依据[11]。但总体而言，国内激光除草技术大多处于实验室阶段，距离产业化应用还有一定差距。')

    add_text('在农业机器人监控系统方面，国外技术较为成熟。美国John Deere公司开发的John Deere Operations Center平台可以实时监控农机的位置、作业状态、作业面积等信息，支持远程诊断和调度管理[12]。欧洲农业机械巨头CNH Industrial开发的MyCNH平台提供类似的农机监控服务，已在全球多个国家和地区部署应用。国内方面，中国农业机械化科学研究院、中联重科等单位也开发了相应的远程监控系统，但在系统功能完备性和用户体验方面与国外先进水平仍有差距[13]。')

    add_text('Web监控系统技术已相当成熟。Vue.js作为当前主流的前端框架，具有轻量、高效、易学习等特点，被广泛应用于各类管理系统的开发中[14]。Spring Boot作为Java生态中最流行的后端开发框架，极大简化了Spring应用的配置和部署过程[15]。WebSocket技术的出现解决了HTTP协议中服务器无法主动推送消息的问题，使得实时监控、在线聊天等应用场景成为可能[16]。这些技术的成熟应用为开发高质量的Web监控系统提供了坚实的技术基础。')

    add_text('综上所述，激光除草技术与Web监控系统技术都已取得长足发展，但专门针对激光除草机器人的Web运行监测系统研究相对较少。本课题将Web技术与激光除草机器人监控需求相结合，设计并实现一套完整的运行监测系统，具有重要的研究价值和实际应用意义。')

    add_heading2('1.3 主要工作内容')

    add_text('本课题主要围绕激光除草机器人运行监测系统展开研究，主要工作内容包括以下几个方面：')

    add_text('（1）需求分析与技术选型：深入分析激光除草机器人监控系统的功能需求与非功能需求，在此基础上进行技术选型，确定前后端开发框架、数据库、通信协议等技术方案。')

    add_text('（2）系统架构设计：设计系统的三层总体架构，包括前端展示层、后端服务层与机载仿真层，设计系统的核心工作流程与数据库表结构。')

    add_text('（3）系统接口设计：设计用户认证、机器人状态管理、检测记录管理、激光控制管理等核心API接口，统一接口规范与数据格式。')

    add_text('（4）前端系统实现：基于Vue 3框架实现系统前端，包括登录页面、仪表盘、状态历史、检测记录、激光控制、轨迹地图等功能页面。')

    add_text('（5）后端服务实现：基于Spring Boot框架实现后端服务，包括用户认证模块、状态管理模块、检测记录模块、激光控制模块等，实现WebSocket实时推送功能。')

    add_text('（6）YOLO模型集成：集成YOLOv8姿态估计模型，实现在线杂草检测推理服务，仅负责模型调用与结果解析，不涉及模型训练与硬件加速工作。')

    add_text('（7）机载仿真端实现：基于Python实现机载仿真端，通过多线程技术模拟机器人的状态上报、检测上报、指令轮询等功能。')

    add_text('（8）系统测试：对系统进行全面的功能测试与性能测试，验证系统的功能完备性与性能指标是否满足设计要求。')

    add_heading2('1.4 论文组织结构')

    add_text('本论文共分为六章，各章内容安排如下：')

    add_text('第1章为绪论，介绍课题的研究背景与意义，分析国内外研究现状，概述主要工作内容与论文组织结构。')

    add_text('第2章为相关技术介绍，介绍系统开发所涉及的前端开发技术、后端开发技术、数据库技术、通信协议技术与目标检测技术。')

    add_text('第3章为系统需求分析，对系统进行概述，详细分析系统的功能需求与非功能需求，并给出系统用例图。')

    add_text('第4章为系统设计，介绍系统的总体架构设计、数据库设计、接口设计、前端页面设计与机载仿真端设计，并给出相应的架构图与类图。')

    add_text('第5章为系统实现与测试，详细介绍系统各模块的实现细节，并进行系统功能测试与性能测试。')

    add_text('第6章为总结与展望，总结课题的主要工作成果，分析系统存在的不足，并对未来的研究方向进行展望。')

    doc.add_page_break()

    # ==================== 第2章 相关技术 ====================
    add_heading1('第2章 相关技术介绍')

    add_heading2('2.1 前端开发技术')

    add_heading3('2.1.1 Vue 3框架')

    add_text('Vue.js是一套用于构建用户界面的渐进式JavaScript框架，由尤雨溪于2014年首次发布。Vue 3是Vue.js的第三个主要版本，于2020年9月正式发布[17]。Vue 3采用Composition API作为新的代码组织方式，相比传统的Options API具有更好的逻辑复用能力与类型推导支持。Vue 3使用Proxy替代Object.defineProperty实现响应式系统，解决了Vue 2中无法检测对象属性添加和数组索引修改的问题，同时大幅提升了响应式系统的性能。')

    add_text('Vue 3还对虚拟DOM进行了重写，引入了PatchFlag、Block Tree等优化技术，通过静态标记和靶向更新大幅提升渲染性能。编译器在编译阶段对模板进行静态分析，标记动态节点，在更新时只处理动态部分，避免了不必要的虚拟DOM对比操作。')

    add_heading3('2.1.2 Vite构建工具')

    add_text('Vite是由Vue.js作者尤雨溪开发的新一代前端构建工具，于2020年首次发布[18]。Vite利用浏览器原生支持的ES Module特性，在开发环境下无需打包，实现秒级启动。传统的构建工具如Webpack在开发模式下需要先打包整个项目，随着项目规模增大，启动时间会越来越长。而Vite在开发模式下将模块的打包工作交给浏览器，服务器只负责按需编译，大大提升了开发体验。')

    add_text('Vite同时支持热更新（HMR）功能，当修改代码时，只会重新编译该模块，而无需重新加载整个页面，热更新速度不会随着项目规模增大而变慢。在生产环境下，Vite使用Rollup进行打包，支持代码分割、Tree Shaking等优化特性。')

    add_heading3('2.1.3 Element Plus组件库')

    add_text('Element Plus是基于Vue 3的UI组件库，是Element UI的Vue 3版本[19]。Element Plus提供了丰富的UI组件，包括按钮、表单、表格、对话框、导航等常用组件，设计遵循一致规范，具有统一的视觉风格。Element Plus完全使用TypeScript编写，提供完整的类型定义，支持按需引入，减少打包体积。Element Plus组件库的组件设计符合企业级应用的交互规范，能够大幅提升前端开发效率。')

    add_heading3('2.1.4 ECharts数据可视化库')

    add_text('ECharts是由百度开发的开源数据可视化库，2018年捐赠给Apache基金会，成为Apache顶级项目[20]。ECharts提供丰富的图表类型，包括折线图、柱状图、散点图、饼图、地图、热力图、关系图等，支持交互操作如缩放、漫游、提示框等。ECharts使用Canvas渲染，性能优秀，支持大数据量展示。ECharts还提供了灵活的配置项，用户可以根据需要自定义图表的样式和行为，满足各类数据可视化需求。')

    add_heading2('2.2 后端开发技术')

    add_heading3('2.2.1 Spring Boot框架')

    add_text('Spring Boot是基于Spring框架的快速开发框架，由Pivotal团队开发，第一个版本发布于2014年[21]。Spring Boot采用"约定优于配置"的理念，提供开箱即用的配置，内嵌Servlet容器（Tomcat、Jetty等），无需部署WAR文件，简化了Spring应用的初始搭建和开发过程。Spring Boot提供丰富的Starter依赖，开发者只需引入对应的Starter即可获得相关技术的自动配置，无需手动配置大量XML文件。Spring Boot还提供了生产就绪特性，如健康检查、指标监控、外部化配置等，便于应用的运维管理。')

    add_heading3('2.2.2 MyBatis Plus')

    add_text('MyBatis Plus是在MyBatis基础上增强的开源持久层框架，旨在简化开发、提高效率[22]。MyBatis Plus提供通用的Mapper与Service，开发者无需编写SQL语句即可实现单表的CRUD操作。MyBatis Plus还提供了分页插件、性能分析插件、乐观锁插件等实用功能，支持代码生成器自动生成实体类、Mapper接口、Service层、Controller层代码。MyBatis Plus只做增强不做改变，引入它不会对现有工程产生影响，兼容所有MyBatis原生特性。')

    add_heading3('2.2.3 JWT认证')

    add_text('JWT（JSON Web Token）是一种开放标准（RFC 7519），用于在各方之间安全地传输信息作为JSON对象[23]。JWT由三部分组成：Header（头部）、Payload（载荷）、Signature（签名）。Header指定签名算法，Payload包含用户信息和声明，Signature用于验证消息在传输过程中未被篡改。JWT具有跨语言支持、自包含、易于传输等优点，是目前最流行的跨域认证解决方案。相比传统的Session认证，JWT无需在服务器端存储会话信息，天然支持分布式部署场景。')

    add_heading2('2.3 数据库技术')

    add_text('MySQL是目前最流行的开源关系型数据库管理系统，由瑞典MySQL AB公司开发，现为Oracle公司产品[24]。MySQL具有开源、免费、性能高、可靠性好等优点，支持标准SQL语言，支持多种存储引擎。InnoDB存储引擎是MySQL 5.5版本后的默认存储引擎，支持事务、外键、行级锁、崩溃恢复等特性，适合大多数业务场景。MySQL支持主从复制、读写分离、分库分表等高可用架构，可满足大规模应用的高并发需求。本系统采用MySQL 8.0存储用户信息、机器人状态、检测记录、激光指令等业务数据。')

    add_heading2('2.4 通信协议技术')

    add_heading3('2.4.1 HTTP与RESTful API')

    add_text('HTTP（HyperText Transfer Protocol）是Web应用中最基础的协议，基于TCP/IP协议，采用请求-响应模型[25]。REST（Representational State Transfer）是一种软件架构风格，由Roy Fielding于2000年提出。RESTful API基于HTTP协议设计，使用HTTP方法表示操作类型：GET表示查询、POST表示创建、PUT表示更新、DELETE表示删除，使用URI表示资源，使用JSON作为数据交换格式。RESTful API具有简洁、易读、易扩展等优点，是目前Web API设计的主流风格。本系统机载仿真端与后端服务之间的通信采用RESTful API风格。')

    add_heading3('2.4.2 WebSocket协议')

    add_text('WebSocket是HTML5提供的一种在单个TCP连接上进行全双工通信的协议，于2011年成为IETF标准RFC 6455[26]。WebSocket解决了HTTP协议中服务器无法主动向客户端推送消息的问题。WebSocket握手采用HTTP协议，客户端发送包含Upgrade头的HTTP请求，服务器响应后建立持久TCP连接，连接建立后双方可以随时发送数据，数据以帧的形式传输，开销小、实时性高。WebSocket适合实时监控、在线聊天、多人协作等需要实时通信的场景。本系统后端与前端之间的实时数据推送采用WebSocket协议。')

    add_heading2('2.5 目标检测技术')

    add_text('YOLO（You Only Look Once）是基于深度学习的目标检测算法，由Joseph Redmon等人于2015年首次提出，将目标检测转化为回归问题，实现端到端检测[27]。YOLO系列算法检测速度快，是工业界应用最广泛的目标检测算法之一。YOLOv8是Ultralytics公司2023年发布的最新版本，提供目标检测、实例分割、姿态估计、图像分类等多种任务支持[28]。YOLOv8在精度和速度上都有较大提升，采用CSPDarknet作为骨干网络，使用PANet进行特征融合，支持多尺度训练。本系统集成YOLOv8姿态估计模型用于杂草检测，获取杂草边界框与茎干关键点信息，为激光除草提供精准定位数据。')

    add_heading2('2.6 本章小结')

    add_text('本章介绍了系统开发涉及的相关技术，包括前端开发技术（Vue 3框架、Vite构建工具、Element Plus组件库、ECharts数据可视化库）、后端开发技术（Spring Boot框架、MyBatis Plus持久层框架、JWT认证技术）、数据库技术（MySQL关系型数据库）、通信协议技术（HTTP协议与RESTful API、WebSocket实时通信协议）和目标检测技术（YOLO系列算法）。这些技术都经过工业界验证，成熟可靠，为系统的设计与实现提供了坚实的技术基础。')

    doc.add_page_break()

    # ==================== 第3章 系统需求分析 ====================
    add_heading1('第3章 系统需求分析')

    add_heading2('3.1 系统概述')

    add_text('激光除草机器人运行监测系统是一套针对激光除草机器人作业过程中的监控管理需求而设计的Web系统。系统旨在实现机器人运行状态的实时监控、杂草检测结果的管理与展示、激光控制指令的下发与执行跟踪等功能，帮助操作人员远程管理激光除草机器人的作业过程。')

    add_text('本系统采用软件仿真的方式，在不依赖实际硬件设备的情况下，验证激光除草机器人监控系统的业务流程与技术可行性。系统通过Python脚本模拟树莓派机载端的行为，实现状态上报、检测上报、指令轮询等功能，与后端服务进行通信交互。本系统仅负责YOLO模型的集成调用，不涉及模型训练和硬件加速工作，激光硬件驱动也不在本系统的开发范围内。')

    add_text('系统采用三层架构设计，包括机载仿真层、后端服务层与前端展示层。机载仿真层负责模拟机器人的运行，上报状态数据与检测结果，轮询并执行指令。后端服务层负责数据的存储与处理，提供API服务与实时推送。前端展示层负责将数据以可视化的形式展示给用户，提供用户交互界面。三层之间通过HTTP协议和WebSocket协议进行通信，实现数据的高效传输与实时推送。')

    add_heading2('3.2 功能需求分析')

    add_text('本系统是基于B/S架构的激光除草机器人运行监测系统，主要包含用户认证、机器人状态监控、检测记录管理、激光控制、仪表盘统计、状态历史查询、轨迹地图展示、YOLO在线检测、操作日志管理等功能模块。')

    add_text('系统用例图如图3-1所示，系统包含两类参与者：系统管理员与普通用户。系统管理员拥有所有功能的操作权限，包括用户管理、系统配置等；普通用户主要拥有查看权限，可以查看机器人状态、检测记录、激光日志等信息，但不能下发激光控制指令。')

    add_caption('图3-1 系统用例图')

    # 添加Mermaid用例图描述
    add_mermaid_diagram('系统用例图', '''graph TD
    Actor[用户] --> Login[登录系统]
    Actor --> Logout[退出登录]
    Actor --> Dashboard[查看仪表盘]
    Actor --> StatusHistory[查询状态历史]
    Actor --> DetectionRecords[查看检测记录]
    Actor --> DetectionDetail[查看检测详情]
    Actor --> LaserControl[激光控制操作]
    Actor --> LaserLogs[查看激光日志]
    Actor --> TrajectoryMap[查看轨迹地图]
    Actor --> YoloDetection[YOLO在线检测]

    subgraph 系统功能
        Login
        Logout
        Dashboard
        StatusHistory
        DetectionRecords
        DetectionDetail
        LaserControl
        LaserLogs
        TrajectoryMap
        YoloDetection
    end''')

    add_heading3('3.2.1 用户认证模块')

    add_text('用户认证模块负责系统的安全认证和权限控制，主要功能包括：')

    add_text('（1）用户登录：用户输入用户名和密码进行登录，系统验证凭据正确性后生成JWT令牌并返回给前端。')

    add_text('（2）用户登出：用户点击登出按钮后，系统将令牌加入黑名单，防止已过期令牌被重复使用。')

    add_text('（3）权限验证：系统所有受保护的接口都需要进行身份验证，后端通过JWT认证过滤器验证请求头中的令牌是否有效。')

    add_text('（4）角色管理：系统支持管理员与普通用户两种角色，不同角色拥有不同的操作权限。')

    add_heading3('3.2.2 机器人状态监控模块')

    add_text('机器人状态监控模块负责机器人状态数据的接收、存储、查询和实时展示，主要功能包括：')

    add_text('（1）状态上报：机载仿真端每5秒上报一次机器人运行状态，包括电量、速度、温度、激光状态、CPU使用率、GPS坐标、IMU数据等。')

    add_text('（2）实时状态展示：前端仪表盘实时展示机器人的当前状态数据，包括核心指标卡片、运行参数趋势图、作业统计等。')

    add_text('（3）WebSocket实时推送：状态数据更新后，后端通过WebSocket将最新状态推送给所有在线的前端客户端。')

    add_text('（4）状态历史查询：支持按时间范围、异常条件等筛选查询历史状态记录，支持分页、排序和导出功能。')

    add_heading3('3.2.3 检测记录管理模块')

    add_text('检测记录管理模块负责杂草检测结果的上报、存储、查询和展示，主要功能包括：')

    add_text('（1）检测结果上报：机载仿真端完成杂草检测后，将检测结果、原始图片、标注图片上报至后端服务。')

    add_text('（2）检测记录查询：支持按时间范围筛选查询检测记录，支持分页显示。')

    add_text('（3）检测详情查看：查看单条检测记录的详细信息，包括标注图片、检测统计、每个检测目标的详细信息。')

    add_text('（4）YOLO在线检测：支持用户上传图片进行在线检测，可调整置信度阈值和类别过滤参数。')

    add_heading3('3.2.4 激光控制模块')

    add_text('激光控制模块负责激光指令的下发、执行、反馈与日志记录，主要功能包括：')

    add_text('（1）指令下发：用户通过前端界面下发激光控制指令，包括开启激光、关闭激光、激光照射、停止、设置功率、自检、复位等。')

    add_text('（2）指令轮询：机载仿真端每1~2秒轮询一次待执行指令，获取指令并模拟执行。')

    add_text('（3）执行反馈：机载仿真端执行完指令后，向服务端上报执行结果，服务端更新指令状态并推送给前端。')

    add_text('（4）激光状态查询：实时查询激光设备的连接状态、运行状态、温度等信息。')

    add_text('（5）操作日志管理：记录所有激光操作的详细日志，支持多条件筛选查询、导出和批量删除。')

    add_heading3('3.2.5 仪表盘与轨迹模块')

    add_text('仪表盘模块提供系统数据的概览展示，轨迹地图模块提供机器人运动轨迹的可视化展示，主要功能包括：')

    add_text('（1）今日统计：展示今日总检测次数、杂草清除总量、作物识别总量、激光作业次数等统计数据。')

    add_text('（2）实时趋势图：展示电量、CPU使用率、温度的实时变化趋势。')

    add_text('（3）目标分布饼图：展示作物与杂草的数量分布情况。')

    add_text('（4）轨迹地图展示：在地图上绘制机器人的运动轨迹，支持标准地图与卫星地图切换。')

    add_text('（5）轨迹回放：支持轨迹的动态回放，可调节播放速度和进度。')

    add_text('（6）轨迹统计：计算并展示总距离、平均速度、作业时长等轨迹统计指标。')

    add_heading2('3.3 非功能需求分析')

    add_heading3('3.3.1 性能需求')

    add_text('性能是衡量系统质量的重要指标，本系统的性能需求主要包括以下几个方面：')

    add_text('（1）响应时间：系统接口的平均响应时间应小于500ms，95%的请求响应时间应小于1s。前端页面的加载时间应小于2s，确保用户操作的流畅性。')

    add_text('（2）并发能力：系统应支持至少50个并发用户同时在线访问，在并发访问情况下系统仍能保持稳定运行，不出现数据错乱或服务崩溃情况。')

    add_text('（3）数据上报频率：机载仿真端状态上报频率为每5秒一次，系统应能稳定处理该频率的数据上报，不出现数据丢失或处理延迟。检测结果上报由于包含图片上传，允许稍长的处理时间。')

    add_text('（4）实时推送延迟：WebSocket实时推送的延迟应小于100ms，确保前端展示的实时性，用户能够及时看到状态变化。')

    add_heading3('3.3.2 安全性需求')

    add_text('安全性是系统设计的重要考虑因素，本系统的安全性需求主要包括以下几个方面：')

    add_text('（1）身份认证：系统所有受保护的接口都需要进行身份认证，防止未授权访问。未登录用户访问受保护接口时应返回401未授权状态码。')

    add_text('（2）密码安全：用户密码应采用加密存储，不允许明文存储密码。推荐使用BCrypt等安全的哈希算法进行密码加密，加盐处理防止彩虹表攻击。')

    add_text('（3）Token安全：认证Token应具有有效期，过期后需要重新登录。Token应采用安全的生成算法，防止被伪造或篡改。系统应支持Token黑名单机制，实现强制登出功能。')

    add_text('（4）数据传输安全：系统应支持HTTPS协议，保证数据在传输过程中的安全性，防止数据被窃听或篡改。')

    add_text('（5）接口防刷：系统应具备接口防刷机制，限制单个用户的接口调用频率，防止恶意调用接口造成系统压力或暴力破解密码。')

    add_heading3('3.3.3 可用性需求')

    add_text('可用性是指系统正常运行的时间比例，本系统的可用性需求主要包括以下几个方面：')

    add_text('（1）系统可用性：系统的可用性应达到99%以上，即全年停机时间不超过87.6小时。系统应具备良好的稳定性，能够长时间稳定运行。')

    add_text('（2）界面友好性：系统界面应简洁直观，操作便捷，符合用户的使用习惯。用户应能在短时间内掌握系统的基本操作，无需复杂培训。界面布局应合理，信息层次清晰，重要信息突出显示。')

    add_text('（3）容错能力：系统应具备良好的容错能力，对于用户的非法输入或异常操作，系统应能给出友好的提示，而不是崩溃或抛出异常信息。系统应进行完善的参数校验，防止非法数据进入系统。')

    add_text('（4）兼容性：系统应兼容主流的浏览器，包括Chrome、Firefox、Safari、Edge等，确保在不同浏览器下都能正常显示与操作。前端页面应采用响应式设计，适配不同尺寸的屏幕。')

    add_heading3('3.3.4 可扩展性需求')

    add_text('可扩展性是指系统应对需求变化的能力，本系统的可扩展性需求主要包括以下几个方面：')

    add_text('（1）功能扩展：系统应采用模块化设计，各模块之间耦合度低，便于后续功能的扩展与升级。新增功能模块不应影响现有模块的正常运行。')

    add_text('（2）硬件扩展：系统设计应考虑后续实际硬件设备的接入，当有真实树莓派与激光设备时，应能方便地接入系统，不需要对系统架构进行大的改动。接口设计应具备通用性和兼容性。')

    add_text('（3）数据扩展：随着系统运行时间的增长，数据量会不断增加，系统应能应对数据量增长带来的挑战，支持数据的高效存储与查询。数据库设计应考虑数据归档和分表策略。')

    add_text('（4）并发扩展：系统应支持水平扩展，当并发用户数增加时，可以通过增加服务器节点的方式提升系统的处理能力。系统架构应支持负载均衡和分布式部署。')

    add_heading2('3.4 本章小结')

    add_text('本章对激光除草机器人运行监测系统进行了全面的需求分析。首先对系统进行了概述，明确了系统的定位与目标，说明本系统采用软件仿真方式验证监控系统可行性，不涉及硬件开发工作。然后详细分析了系统的功能需求，包括用户认证模块、机器人状态监控模块、检测记录管理模块、激光控制模块、仪表盘与轨迹模块等五个核心模块的具体功能需求，并给出了系统用例图。最后分析了系统的非功能需求，包括性能需求、安全性需求、可用性需求与可扩展性需求。需求分析是系统设计与实现的基础，本章的内容为后续的系统设计提供了明确的依据和指导方向。')

    doc.add_page_break()

    # ==================== 第4章 系统设计 ====================
    add_heading1('第4章 系统设计')

    add_heading2('4.1 系统总体架构设计')

    add_heading3('4.1.1 三层架构设计')

    add_text('本系统采用经典的三层B/S架构设计，从上到下依次为前端展示层、后端服务层与机载仿真层。三层架构具有职责清晰、耦合度低、可扩展性强等优点，便于系统的开发、维护与升级。')

    add_text('系统总体架构图如图4-1所示。')

    add_caption('图4-1 系统总体架构图')

    add_mermaid_diagram('系统总体架构图', '''graph TB
    subgraph 前端展示层
        Login[登录页面<br>Vue 3 + Element Plus]
        Dashboard[仪表盘<br>ECharts可视化]
        Status[状态历史]
        Detection[检测记录]
        Laser[激光控制]
        Trajectory[轨迹地图]
        YoloPage[YOLO检测]
    end

    subgraph 后端服务层
        Controller[Controller层<br>API接口]
        Service[Service层<br>业务逻辑]
        Mapper[Mapper层<br>数据持久化]
        WebSocket[WebSocket服务<br>实时推送]
    end

    subgraph 机载仿真层
        StateThread[状态上报线程<br>每5秒]
        DetectThread[检测上报线程<br>每5-10秒]
        CmdThread[指令轮询线程<br>每1-2秒]
        YoloService[YOLO推理服务]
    end

    subgraph 数据存储层
        MySQL[(MySQL数据库<br>用户/状态/检测/指令)]
        FileStorage[本地文件存储<br>检测图片]
    end

    Login -->|HTTP/REST| Controller
    Dashboard -->|WebSocket| WebSocket
    Status -->|HTTP/REST| Controller
    Detection -->|HTTP/REST| Controller
    Laser -->|HTTP/REST| Controller
    Trajectory -->|HTTP/REST| Controller
    YoloPage -->|HTTP/REST| Controller

    Controller --> Service
    Service --> Mapper
    Mapper --> MySQL
    Service --> FileStorage

    StateThread -->|POST /robot/status| Controller
    DetectThread -->|POST /detection/report| Controller
    CmdThread -->|GET /robot/commands/pending| Controller

    DetectThread -->|调用| YoloService

    WebSocket -.->|实时推送| Dashboard
    WebSocket -.->|实时推送| Status
    WebSocket -.->|实时推送| Laser''')

    add_text('各层的职责如下：')

    add_text('（1）前端展示层：负责与用户进行交互，展示系统数据，接收用户操作。前端采用Vue 3框架开发，使用Element Plus作为UI组件库，使用ECharts进行数据可视化。前端与后端通过HTTP协议进行API调用，通过WebSocket协议接收实时数据推送。前端包含登录页面、仪表盘页面、状态历史页面、检测记录页面、激光控制页面、轨迹地图页面、YOLO在线检测页面等功能页面。')

    add_text('（2）后端服务层：负责业务逻辑处理、数据存储与API服务提供。后端采用Spring Boot框架开发，提供RESTful API接口与WebSocket服务。后端采用经典的三层架构：Controller层接收HTTP请求，Service层处理业务逻辑，Mapper层负责数据库操作。后端包含用户认证、状态管理、检测记录、激光控制、统计分析等核心业务模块，使用MySQL数据库进行数据持久化存储，使用本地文件系统存储检测图片。')

    add_text('（3）机载仿真层：负责模拟树莓派机载端的行为，与后端服务进行通信。机载仿真端采用Python开发，使用多线程技术实现状态上报、检测上报、指令轮询等并发任务。状态上报线程每5秒上报一次机器人运行状态，检测上报线程每5-10秒进行一次杂草检测并上报结果，指令轮询线程每1-2秒查询一次待执行激光指令。机载仿真端还集成了YOLO推理服务，用于模拟杂草检测过程。')

    add_heading3('4.1.2 系统工作流程设计')

    add_text('系统的核心工作流程描述了激光除草机器人从登录开始到完成作业的完整过程。系统的核心工作流程如下：')

    add_text('（1）用户通过前端页面输入用户名与密码进行登录，前端将登录请求发送至后端，后端验证凭据正确性后生成JWT令牌并返回给前端，前端将令牌保存到本地存储中。')

    add_text('（2）前端建立与后端的WebSocket连接，用于接收实时数据推送。连接建立后，后端会推送最新的机器人状态数据。')

    add_text('（3）机载仿真端登录后端获取认证令牌，建立通信会话。机载仿真端启动三个工作线程：状态上报线程、检测上报线程、指令轮询线程。')

    add_text('（4）机载仿真端每5秒上报一次机器人运行状态，后端保存状态数据到数据库，并通过WebSocket将最新状态推送给所有在线的前端客户端。')

    add_text('（5）机载仿真端完成杂草检测后，将检测结果、原始图片、标注图片上报至后端，后端保存检测记录并通过WebSocket推送给前端，更新仪表盘统计数据。')

    add_text('（6）用户通过前端激光控制页面下发激光控制指令，后端保存指令到数据库并返回成功响应。')

    add_text('（7）机载仿真端每1~2秒轮询一次待执行指令，获取到新指令后模拟执行过程，根据指令类型设置不同的执行时长。')

    add_text('（8）机载仿真端执行完指令后，向服务端发送执行反馈，服务端更新指令状态并通过WebSocket将执行结果推送给前端，前端更新激光控制页面显示。')

    add_heading2('4.2 数据库设计')

    add_heading3('4.2.1 数据库概念设计')

    add_text('根据系统需求分析，本系统的核心实体包括用户、机器人信息、机器人状态、检测记录、激光指令、激光操作日志、JWT黑名单等。各实体之间的关系如下：用户与机器人为一对多关系（一个用户可管理多个机器人）；机器人与机器人状态为一对多关系（一个机器人有多条状态记录）；机器人与检测记录为一对多关系（一个机器人产生多条检测记录）；机器人与激光指令为一对多关系（一个机器人执行多条激光指令）。')

    add_heading3('4.2.2 数据库表结构设计')

    add_text('根据数据库概念设计，设计系统的数据库表结构。本系统共包含8张核心数据表，各表的详细结构如下：')

    add_text('（1）用户表（sys_user）：存储系统用户信息，包括用户ID、用户名、密码哈希、角色、创建时间、更新时间等字段。用户ID为主键，用户名为唯一索引。')

    add_text('（2）机器人信息表（robot_info）：存储机器人基本信息，包括机器人ID、机器人编号、名称、型号、状态、创建时间等字段。')

    add_text('（3）机器人状态表（robot_status）：存储机器人运行状态历史，包括状态ID、机器人ID、电量、速度、温度、激光状态、CPU使用率、经度、纬度、IMU数据、上报时间等字段。上报时间建立索引以提高查询效率。')

    add_text('（4）检测记录表（detection_record）：存储杂草检测结果，包括记录ID、原始图片路径、标注图片路径、杂草数量、作物数量、推理耗时、检测时间等字段。检测时间建立索引。')

    add_text('（5）检测目标表（detection_item）：存储单次检测中的每个目标详情，与检测记录表为一对多关系，包括目标ID、检测记录ID、目标类别、置信度、边界框坐标、关键点坐标、深度、三维坐标等字段。')

    add_text('（6）激光指令表（laser_command）：存储激光控制指令，包括指令ID、指令编号、动作类型、参数、状态、创建时间、执行时间、完成时间等字段。')

    add_text('（7）激光操作日志表（laser_operation_log）：存储激光操作执行日志，包括日志ID、指令ID、动作、目标坐标、时长、功率、执行结果、执行消息、创建时间等字段。')

    add_text('（8）JWT黑名单表（jwt_blacklist）：存储已失效的JWT令牌，实现强制登出功能，包括令牌ID、令牌值、过期时间等字段。')

    add_heading2('4.3 主要类图设计')

    add_text('类图是描述系统中类的静态结构的重要工具，通过类图可以清晰地展示系统中各个类之间的关系，包括继承、关联、依赖等。本系统主要包含系统安全与认证、机器人状态与轨迹、激光设备控制、杂草检测与管理四大核心模块。')

    add_heading3('4.3.1 系统安全与认证模块类图')

    add_text('系统安全与认证模块采用经典的三层架构设计：UserController接收前端登录、登出请求；UserService处理业务逻辑，包括密码验证、JWT令牌生成与验证；UserMapper负责数据库的读写操作；User实体类与数据库中的sys_user表相对应，包含用户ID、用户名、密码哈希、角色等字段；JwtBlacklistService管理失效的JWT令牌，实现登出功能。')

    add_heading3('4.3.2 机器人状态与轨迹模块类图')

    add_text('机器人状态与轨迹模块整合了状态监控、历史查询和轨迹分析功能。RobotStatus实体类包含机器人运行状态的所有关键指标。RobotStatusController提供状态上报、最新状态查询、历史状态查询与轨迹查询四个核心接口。RobotStatusService负责状态数据的处理、轨迹距离计算和统计分析。WebSocketHandler负责将状态更新实时推送给前端客户端。TrajectoryPoint用于封装轨迹点信息，包括经度、纬度、时间、状态信息等。')

    add_heading3('4.3.3 激光设备控制模块类图')

    add_text('激光设备控制模块包含LaserCommand与LaserOperationLog两个核心实体，分别对应激光指令与激光执行日志。LaserController负责接收前端激光控制请求，提供指令下发、状态查询、日志查询等接口。LaserControlService负责业务逻辑处理，包括指令创建、执行结果更新、日志管理等。LaserCommandMapper与LaserOperationLogMapper负责各自实体的数据库操作。模块支持指令的下发、状态查询、执行反馈，以及操作日志的查询、导出和批量删除功能。')

    add_heading3('4.3.4 杂草检测与管理模块类图')

    add_text('杂草检测与管理模块整合了检测结果管理和在线推理功能。DetectionController负责接收检测上报、记录查询、在线检测请求。DetectionService负责业务逻辑处理，包括检测结果解析、图片存储、记录查询等。DetectionRecordMapper与DetectionItemMapper负责检测记录和检测目标的数据库操作。YoloInferenceService用于调用YOLO模型进行推理，通过执行Python脚本调用Ultralytics YOLOv8模型。FileStorageService用于文件存储管理，负责原始图片和标注图片的存储与访问。PredictionView是推理结果的视图对象，Prediction类封装单个检测目标的信息，包括类别、置信度、边界框、关键点、深度、三维坐标等。')

    add_heading2('4.4 功能模块详细设计')

    add_heading3('4.4.1 系统安全与认证模块')

    add_text('系统安全与认证模块是系统访问控制的基础，负责验证用户身份并管理用户的登录状态。用户登录认证流程如下：用户在登录页面输入用户名和密码，点击登录按钮；前端将用户名和密码发送至后端登录接口；后端接收请求后根据用户名查询数据库中的用户信息；后端使用BCrypt算法验证密码是否正确；密码验证通过后，后端生成JWT令牌，设置有效期为24小时；后端返回登录成功响应，包含令牌和用户基本信息；前端将令牌保存到本地存储localStorage中，保存用户信息到Pinia状态管理；前端跳转到系统首页。')

    add_text('该模块采用BCrypt算法进行密码加密存储，避免明文密码泄露风险。BCrypt算法自动加盐，每次加密相同密码得到的哈希值都不同，有效防止彩虹表攻击。使用JWT令牌管理用户登录状态，Token默认有效期为24小时，超时后需要重新登录。系统支持管理员与普通用户两种角色，不同角色拥有不同的操作权限，通过Spring Security实现基于角色的访问控制。用户登出时，系统将令牌加入黑名单，防止已过期令牌被重复使用。')

    add_heading3('4.4.2 机器人状态与轨迹模块')

    add_text('机器人状态与轨迹模块负责实时采集机器人运行状态、历史状态查询以及运动轨迹可视化展示，是系统监控功能的核心模块。状态数据采集与实时监控流程如下：机载仿真端的状态上报线程启动，每5秒执行一次上报任务；线程生成模拟的状态数据，包括电量、温度、CPU使用率、GPS坐标、速度、IMU数据等，各参数在合理范围内随机变化；线程将状态数据封装为JSON格式，携带认证Token调用POST /api/v1/robot/status接口；后端RobotStatusController接收请求，验证Token有效性；RobotStatusService将状态数据保存到robot_status表；后端通过WebSocketHandler将最新状态数据推送给所有在线的前端客户端；前端收到WebSocket消息后，更新仪表盘页面的状态卡片和图表数据。')

    add_text('历史状态查询支持多维度筛选，包括时间范围、激光状态、异常条件（电量过低、温度过高、CPU过载）以及经纬度关键词搜索。系统采用分页查询方式，每页默认20条记录，查询结果支持按上报时间倒序排列，支持导出CSV格式文件。轨迹地图功能支持标准地图与卫星视图切换，轨迹点超过1000个时自动进行抽稀处理以保证前端渲染性能。轨迹回放支持播放/暂停、多倍速调节（0.5x~4x）和进度拖动。系统采用Haversine公式计算球面距离，统计指标包括总距离、平均速度、最大速度、作业时长等。')

    add_heading3('4.4.3 激光设备控制模块')

    add_text('激光设备控制模块负责激光指令的下发、轮询执行、结果反馈和操作日志管理，实现了从指令发送到执行记录的完整闭环。激光指令执行与日志管理流程如下：用户在激光控制页面点击控制按钮或设置参数后点击执行；前端发送POST请求到激光指令下发接口；后端LaserController接收请求，验证用户权限；LaserControlService创建激光指令记录，生成唯一指令ID，设置初始状态为"待执行"，保存到数据库；后端返回成功响应给前端；机载仿真端的指令轮询线程每1-2秒调用GET /api/v1/robot/commands/pending接口查询待执行指令；后端返回状态为"待执行"的指令列表；机载仿真端获取到指令后，根据指令类型模拟执行过程，设置不同的执行时长；指令执行完成后，机载仿真端调用POST /api/v1/robot/laser/feedback接口上报执行结果，同时调用PUT /api/v1/robot/commands/{id}/ack接口确认指令已执行；后端更新指令状态，创建激光操作日志记录；后端通过WebSocket将指令执行结果推送给前端；前端更新激光控制页面的指令执行状态和日志列表。')

    add_text('系统支持多种指令类型：ENABLE（设备上电使能）、DISABLE（设备断电关闭）、FIRE（对指定坐标执行激光照射）、STOP（立即停止当前照射）、SET_POWER（设置激光输出功率）、AIM（瞄准指定坐标但不发射）、SELF_TEST（执行设备自检）、RESET（设备复位）。每条指令的执行结果都会被完整记录，包括目标坐标、深度、时长、功率、执行结果、说明信息和创建时间。操作日志查询支持按时间范围、指令类型、执行结果等多条件组合筛选，提供快捷时间选项（今天、昨天、近7天、近30天）和数值范围筛选。执行结果字段使用不同颜色标签展示，支持批量导出CSV和批量删除操作。')

    add_heading3('4.4.4 杂草检测与管理模块')

    add_text('杂草检测与管理模块集成了YOLO在线检测、检测结果上报和检测记录管理三大功能，是系统核心的杂草识别与数据管理模块。检测结果上报流程如下：机载仿真端的检测上报线程每5-10秒执行一次检测任务；线程从测试图片库中随机选择一张图片；线程调用后端YOLO检测接口，上传图片；后端YoloInferenceService接收图片，调用Python推理脚本执行YOLOv8姿态估计模型推理；推理完成后，后端解析JSON格式的推理结果，生成标注图片，返回推理结果；机载仿真端收到推理结果后，调用POST /api/v1/detection/report接口上报检测结果，包含原始图片、标注图片和检测目标列表；后端DetectionService接收上报数据，保存图片到本地文件系统，保存检测记录和检测目标到数据库；后端通过WebSocket推送检测结果通知；前端收到通知后更新仪表盘的检测统计数据。')

    add_text('YOLO在线检测流程允许用户上传JPEG或PNG格式图片，设置检测置信度阈值（默认0.25，范围0.01~1.0）和类别过滤参数（全部、仅杂草、仅作物）。推理服务使用Python脚本调用Ultralytics YOLOv8姿态估计模型，输出检测结果包括目标类别ID、类别名称、置信度、边界框坐标和茎部关键点坐标。前端提供原图与标注图片对比展示，支持图片缩放查看，以及检测统计信息和目标列表展示。目标列表表格展示每个检测目标的详细信息，包括类别、置信度、像素坐标、关键点信息、深度、三维坐标等，支持按类别筛选。')

    add_heading2('4.5 机载仿真端设计')

    add_text('机载仿真端是系统的重要组成部分，负责模拟真实机器人的运行行为，包括状态上报、杂草检测、指令执行等功能。通过软件仿真的方式，可以在不依赖真实硬件设备的情况下，完整验证系统的功能与性能。本系统的机载仿真端采用Python语言开发，利用Python丰富的第三方库和简洁的语法，快速实现模拟功能。')

    add_heading3('4.5.1 多线程架构设计')

    add_text('机载仿真端采用多线程架构设计，使用三个独立的工作线程分别处理状态上报、检测上报与指令轮询任务，模拟真实系统中的并行操作。主线程负责初始化工作和资源管理。机载仿真端的多线程架构图如图4-2所示。')

    add_caption('图4-2 机载仿真端多线程架构图')

    add_mermaid_diagram('机载仿真端多线程架构图', '''graph TB
    Main[主线程<br>MainThread] -->|创建启动| State[状态上报线程<br>Thread-Status<br>每5秒执行]
    Main -->|创建启动| Detect[检测上报线程<br>Thread-Detect<br>每5-10秒执行]
    Main -->|创建启动| Cmd[指令轮询线程<br>Thread-Command<br>每1-2秒执行]

    Main --> Config[配置加载器<br>ConfigLoader]
    Main --> Http[HTTP会话管理器<br>HttpSession]
    Main --> Queue[线程安全队列<br>ThreadSafeQueue]
    Main --> Event[事件标志<br>EventFlag]

    State -->|上报状态| Http
    State -->|状态数据| Queue

    Detect -->|调用YOLO| Yolo[YOLO推理服务<br>YoloInference]
    Detect -->|上报检测结果| Http
    Detect -->|检测数据| Queue

    Cmd -->|查询待执行指令| Http
    Cmd -->|上报执行反馈| Http

    Yolo -->|推理结果| Detect

    style Main fill:#f9f,stroke:#333,stroke-width:2px
    style State fill:#9f9,stroke:#333
    style Detect fill:#99f,stroke:#333
    style Cmd fill:#ff9,stroke:#333''')

    add_text('各线程的职责如下：')

    add_text('（1）主线程：负责初始化仿真器，加载配置参数（服务端地址、认证信息、上报间隔等），建立与后端服务的HTTP会话，启动其他三个工作线程，管理程序生命周期，处理退出信号与资源清理。主线程维护线程安全队列用于线程间数据通信，使用事件标志实现线程间的同步与停止控制。程序退出时，主线程设置停止事件，通知各工作线程安全退出，等待所有工作线程结束后释放资源并退出程序。')

    add_text('（2）状态上报线程：负责定时采集机器人状态数据并上报至服务端，每5秒执行一次。状态数据包括电量、温度、CPU使用率、GPS坐标、速度、IMU数据（加速度、角速度、俯仰角、横滚角）等。每次上报时在合理范围内随机变化这些参数，模拟真实机器人的状态变化。电量模拟缓慢下降趋势，温度和CPU使用率在正常范围内波动，GPS坐标缓慢移动模拟机器人前进，IMU数据加入小幅噪声模拟传感器真实输出。状态数据通过HTTP POST /api/v1/robot/status接口上报。')

    add_text('（3）检测上报线程：负责模拟杂草检测过程，每5-10秒执行一次。线程从样本图片库中随机选择一张测试图片，调用YOLO推理服务执行杂草检测，推理服务返回检测框、置信度、类别、关键点坐标等结果。线程根据推理结果在原始图片上绘制检测框和关键点，生成标注图片。然后调用POST /api/v1/detection/report接口上报检测结果，包括原始图片、标注图片、杂草数量、作物数量、推理耗时以及每个检测目标的详细信息。')

    add_text('（4）指令轮询线程：负责定时轮询待执行的激光指令，获取并执行指令，轮询间隔在1秒到2秒之间随机。线程调用GET /api/v1/robot/commands/pending接口获取待执行指令列表。获取到指令后模拟执行过程，根据指令类型设置不同的执行时长：FIRE指令根据照射时长参数设置，SELF_TEST指令模拟较长时间的自检过程，其他指令快速完成。执行完成后根据指令类型生成相应的执行结果和消息。执行完成后，调用POST /api/v1/robot/laser/feedback接口上报执行结果，同时调用PUT /api/v1/robot/commands/{id}/ack接口确认指令已执行。')

    add_heading3('4.5.2 通信机制设计')

    add_text('机载仿真端与后端服务之间的通信采用HTTP协议。由于激光控制模块为预留模块，对实时性要求不高，采用轮询方式获取指令可以简化实现，降低开发复杂度。通信机制架构如图4-3所示。')

    add_caption('图4-3 机载仿真端通信机制架构图')

    add_mermaid_diagram('通信机制架构图', '''graph LR
    subgraph 机载仿真端
        A1[状态上报线程]
        A2[检测上报线程]
        A3[指令轮询线程]
    end

    subgraph HTTP通信层
        B1[POST /robot/status]
        B2[POST /detection/report]
        B3[GET /robot/commands/pending]
        B4[POST /robot/laser/feedback]
        B5[PUT /robot/commands/ack]
    end

    subgraph 后端服务
        C1[RobotStatusController]
        C2[DetectionController]
        C3[LaserController]
        C4[CommandController]
    end

    subgraph 数据存储
        D1[(robot_status表)]
        D2[(detection_record表)]
        D3[(laser_command表)]
        D4[(laser_operation_log表)]
    end

    A1 --> B1
    B1 --> C1
    C1 --> D1

    A2 --> B2
    B2 --> C2
    C2 --> D2

    A3 --> B3
    B3 --> C4
    C4 --> D3

    A3 --> B4
    B4 --> C3
    C3 --> D4

    A3 --> B5
    B5 --> C4
    C4 --> D3

    style A1 fill:#9f9,stroke:#333
    style A2 fill:#99f,stroke:#333
    style A3 fill:#ff9,stroke:#333''')

    add_text('上行通信（机载仿真端 → 后端服务）：机载仿真端通过调用HTTP POST接口主动上报数据。上行通信的数据流包括：')

    add_text('（1）运行状态上报：调用POST /api/v1/robot/status接口，每5秒上报一次，上报内容包括电量、温度、CPU使用率、GPS坐标、速度、IMU数据（加速度、角速度、俯仰角、横滚角）、上报时间戳等。请求头携带JWT认证令牌，请求体采用JSON格式。')

    add_text('（2）检测结果上报：调用POST /api/v1/detection/report接口，在每次检测完成时上报。采用multipart/form-data格式上传原始图片和标注图片，同时通过JSON格式的result字段传递检测结果，包括杂草数量、作物数量、推理耗时、每个检测目标的类别、置信度、边界框、关键点、深度、三维坐标等信息。')

    add_text('（3）激光执行反馈上报：调用POST /api/v1/robot/laser/feedback接口，在执行完激光指令后上报。上报内容包括指令ID、动作类型、执行结果（SUCCESS/FAILED/TIMEOUT）、执行消息、时间戳等。')

    add_text('下行通信（后端服务 → 机载仿真端）：机载仿真端通过定时调用HTTP GET接口轮询待执行指令。下行通信的数据流包括：')

    add_text('（1）待执行指令获取：调用GET /api/v1/robot/commands/pending接口，每1~2秒轮询一次，获取待执行的激光指令列表。返回结果按创建时间升序排列，保证先创建的指令先被执行。返回的指令信息包括指令ID、指令编号、动作类型、参数JSON、创建时间等。')

    add_text('（2）指令执行确认：调用PUT /api/v1/robot/commands/{id}/ack接口，在执行完指令后调用，标记指令为已执行状态，防止指令被重复执行。请求体包含执行结果和执行消息。')

    add_heading2('4.6 本章小结')

    add_text('本章详细介绍了激光除草机器人运行监测系统的设计方案。首先介绍了系统的总体三层B/S架构设计，包括前端展示层、后端服务层与机载仿真层，并设计了系统的核心工作流程，给出了总体架构图。然后进行了数据库设计，包括数据库概念设计与8张核心数据表的表结构设计。接着进行了主要类图设计，根据实际代码结构，给出了用户认证、机器人状态、检测记录、激光控制四个核心模块的类图结构说明。之后进行了功能模块详细设计，使用流程图详细描述了用户认证、机器人状态监控、检测记录管理、激光控制四个模块的工作流程。最后进行了机载仿真端设计，包括多线程架构设计与通信机制设计，给出了相应的架构图。本章的设计方案为后续的系统实现提供了明确的指导和依据。')

    doc.add_page_break()

    # ==================== 第5章 系统实现与测试 ====================
    add_heading1('第5章 系统实现与测试')

    add_heading2('5.1 开发环境搭建')

    add_heading3('5.1.1 后端开发环境')

    add_text('后端开发使用Java 17语言，开发工具为IntelliJ IDEA 2023，构建工具为Maven 3.9。后端主要依赖包括：Spring Boot 3.2.x Web框架，提供RESTful API支持；Spring WebSocket实时通信支持；MyBatis Plus 3.5.x数据持久化增强工具；MySQL 8.0.x数据库驱动；JWT 0.11.5令牌生成与验证工具；SpringDoc OpenAPI文档生成工具；Lombok简化Java代码工具。')

    add_text('后端配置文件application.yml包含：服务器端口配置（默认8080）、数据库连接配置（MySQL连接URL、用户名、密码）、文件上传大小限制（单文件最大50MB）、MyBatis Plus配置（ mapper位置、主键策略）、JWT密钥和有效期配置（24小时）、文件存储根路径配置、YOLO模型相关配置（Python可执行文件路径、推理脚本路径、模型权重路径、输出目录等）。')

    add_heading3('5.1.2 前端开发环境')

    add_text('前端开发使用Node.js 18.x LTS环境，开发工具为Visual Studio Code，构建工具为Vite 5.x。前端主要依赖包括：Vue 3.4.x核心框架、Vue Router 4.x路由管理、Pinia 2.x状态管理、Element Plus 2.5.x UI组件库、ECharts 5.4.x图表可视化、Axios 1.6.x HTTP请求库。')

    add_text('前端项目结构清晰：src/api目录存放API接口调用函数；src/assets目录存放静态资源文件（图片、样式等）；src/components目录存放通用Vue组件；src/layout目录存放布局组件（侧边栏、顶部导航等）；src/router目录存放路由配置；src/store目录存放Pinia状态管理；src/utils目录存放工具函数（axios封装、WebSocket封装等）；src/views目录存放页面组件；App.vue为根组件；main.js为入口文件。')

    add_heading2('5.2 模块实现')

    add_heading3('5.2.1 用户认证与授权模块')

    add_text('用户认证与授权模块负责系统的安全认证和权限控制，采用Spring Security + JWT的方式实现。用户密码使用BCrypt算法加密存储，确保用户信息安全。登录界面采用居中布局，包含系统Logo、用户名输入框、密码输入框和登录按钮，整体风格简洁现代。用户输入正确的用户名和密码后点击登录按钮，前端通过axios发送POST请求到/api/v1/auth/login接口，后端验证凭据正确性后生成JWT令牌并返回给前端，令牌有效期为24小时。前端将令牌保存到本地存储localStorage中，将用户信息保存到Pinia的user store中，并自动跳转到首页。')

    add_text('前端在后续请求的Authorization头中携带该令牌，格式为"Bearer {token}"。后端通过JWT认证过滤器进行验证，解析令牌获取用户ID，查询用户信息设置到SecurityContext中。若令牌过期、签名错误或无效，系统会返回401未授权状态码，前端拦截401响应自动清除本地存储的令牌并跳转到登录页面。')

    add_text('登出功能实现了令牌黑名单管理，用户点击顶部导航栏的退出登录按钮后，前端发送POST请求到/api/v1/auth/logout接口，后端将令牌加入jwt_blacklist表，防止已过期令牌被重复使用，同时前端清除本地存储的令牌和Pinia中的用户信息，跳转到登录页面。系统支持管理员（ADMIN）与普通用户（USER）两种角色，不同角色拥有不同的菜单权限和操作权限。')

    add_heading3('5.2.2 机器人状态监测模块')

    add_text('机器人状态监测模块负责机器人状态数据的接收、存储、查询和实时展示，是系统的核心功能之一。状态上报功能实现了机器人状态数据的接收和存储，树莓派每5秒调用POST /api/v1/robot/status接口上报一次状态数据，包括电量、速度、温度、激光状态、CPU使用率、GPS坐标（经度、纬度）、IMU数据（加速度x/y/z、角速度x/y/z、俯仰角、横滚角）等。后端接收数据后保存到robot_status表，同时更新机器人当前状态和激光连接状态，并通过WebSocket实时广播状态更新消息给所有在线客户端。')

    add_text('仪表盘页面是状态监测的核心界面，采用玻璃态卡片设计风格，顶部为系统标题和导航栏，左侧为功能菜单，右侧为主内容区。仪表盘主要包含以下部分：')

    add_text('（1）核心实时状态卡片：以2×4网格布局展示8个核心指标卡片，包括机器人状态、当前速度、GPS坐标、激光状态、剩余电量、CPU使用率、机身温度、作业时长。机器人状态分为离线、待机、作业中、故障等，用不同颜色标识（离线灰色、待机蓝色、作业中绿色、故障红色）。实时状态卡片右上角有绿色闪烁的"实时"标识，表明数据正在实时更新。电量卡片带有环形进度条，直观展示剩余电量百分比。')

    add_text('（2）设备运行参数趋势图：使用ECharts绘制三个折线图，分别展示最近20条记录的电量、CPU使用率、温度变化趋势。图表带有时间横轴和数值纵轴，支持鼠标悬停显示详细数值，数据随WebSocket状态上报自动更新，曲线平滑过渡。')

    add_text('（3）今日检测目标分布：使用ECharts饼图展示作物和杂草的数量分布，饼图带有图例和百分比标签，帮助用户快速了解当天的检测情况和杂草占比。')

    add_text('（4）今日作业统计：以四个数字卡片展示今日总检测次数、杂草清除总量、作物识别总量、激光作业次数等统计数据，数据随检测结果上报实时更新。')

    add_text('（5）IMU姿态信息卡片：展示机器人的IMU传感器数据，包括三轴加速度、三轴角速度、俯仰角、横滚角等信息，采用数值和小仪表结合的展示方式。')

    add_text('状态历史页面提供机器人状态历史的查询、筛选和导出功能：')

    add_text('（1）时间范围筛选器：顶部放置日期范围选择器，支持快捷选择（今天、昨天、近7天、近30天）和自定义时间范围，用户点击快捷按钮或选择起止日期后自动刷新数据。')

    add_text('（2）高级筛选面板：点击展开按钮显示高级筛选选项，提供激光状态下拉选择、电量最大值输入、温度最小值输入、CPU使用率最小值输入等异常条件筛选。用户可以组合多个筛选条件快速定位异常状态记录。')

    add_text('（3）数据表格：使用Element Plus的el-table组件展示状态历史记录，包括上报时间、电量、速度、温度、激光状态、CPU使用率、经度、纬度等列。表格支持分页显示，每页显示20条记录，底部分页器支持切换每页条数和跳转到指定页。')

    add_text('（4）排序功能：点击表格列头可以按该字段升序或降序排列，支持多列组合排序。时间列默认倒序排列，最新的记录显示在最前面。')

    add_text('（5）统计信息：表格上方显示总记录数、最高温度、平均温度、最低电量、平均CPU、最高CPU、平均速度、激光开启次数等统计数据，数据颜色根据数值大小动态变化，异常值标红显示。')

    add_text('（6）导出功能：顶部工具栏提供导出按钮，支持导出当前筛选结果为CSV文件，导出的文件包含所有字段，可直接用Excel打开进行后续分析。')

    add_text('（7）快捷查看按钮：顶部提供三个快捷按钮，分别是"查看低电量"（电量<20%）、"查看高温"（温度>40℃）、"查看高CPU"（CPU>80%），点击后自动设置相应筛选条件并刷新数据，方便用户快速定位异常数据。')

    add_heading3('5.2.3 轨迹管理模块')

    add_text('轨迹管理模块负责机器人运动轨迹的记录、查询和可视化展示，为用户提供机器人作业路径的直观呈现。轨迹地图页面使用高德地图JavaScript API实现机器人运动轨迹的可视化展示：')

    add_text('（1）时间范围选择：页面顶部放置时间范围选择器，用户可以选择需要查询轨迹的时间段，支持快捷选择和自定义时间范围，选择后自动加载对应时间段的轨迹数据。')

    add_text('（2）地图切换控件：地图右上角提供地图模式切换按钮，支持标准地图、卫星影像地图、卫星+路网地图三种模式，用户可以根据需要切换不同的地图视图。')

    add_text('（3）轨迹显示：在地图上绘制机器人的运动轨迹线，轨迹线采用渐变色，从绿色到红色表示时间从早到晚，直观展示机器人的移动方向和路径。轨迹线宽度适中，带有半透明效果保证地图底图可见。')

    add_text('（4）轨迹点信息：地图上显示轨迹采样点，点击轨迹点弹出信息窗口，显示该点的详细信息，包括时间、电量、速度、温度、激光状态、CPU使用率等。')

    add_text('（5）轨迹统计信息：页面左侧显示轨迹统计面板，展示总行驶距离、平均速度、最大速度、作业时长、起始位置坐标、结束位置坐标等信息。总距离采用Haversine公式计算相邻两点的球面距离累加得到。')

    add_text('（6）轨迹回放控制：页面底部放置轨迹回放控制面板，提供播放/暂停按钮、速度调节滑块（0.5x、1x、2x、5x四档）、进度条拖动功能。回放时地图上动态显示机器人位置标记沿轨迹移动，同时更新左侧当前点的信息，播放速度实时生效。')

    add_heading3('5.2.4 激光设备控制模块')

    add_text('激光设备控制模块负责激光设备的指令下发、执行反馈和日志记录，采用轮询机制实现指令下发，简化了系统实现复杂度。激光控制页面采用上下布局，上方为设备状态和控制区，下方为执行日志区：')

    add_text('（1）设备状态卡片：页面左上角显示设备状态卡片，包含连接状态指示灯、当前状态文本、当前功率值、上次操作时间等信息。连接状态分为已连接（绿色圆点）和未连接（灰色圆点），当前状态分为待机（蓝色）、作业中（绿色）、故障（红色）等，状态文本带有对应颜色的标签。')

    add_text('（2）指令控制面板：页面右上角为指令按钮区，以网格布局放置各种控制按钮，包括开启激光、关闭激光、开始照射、停止照射、设备自检、设备复位、设置功率等按钮。按钮带有图标和文字，点击按钮后按钮显示加载状态，防止重复点击。')

    add_text('（3）坐标与参数设置区域：页面中部为参数设置区，包含目标X坐标输入框、目标Y坐标输入框、深度值输入框、照射时长滑块（100-5000毫秒）、功率值输入框等参数，用户可以根据需要设置精细的控制参数，输入框带有合理的默认值和数值范围限制。')

    add_text('（4）执行日志区域：页面下方为执行日志表格，显示最近的指令执行记录，包括时间、指令ID、动作类型、执行结果、执行消息等信息，方便用户查看指令执行历史。执行结果列使用不同颜色的标签显示，成功为绿色，失败为红色，超时为橙色。')

    add_text('指令下发功能实现了控制指令的创建和存储。用户点击控制按钮后，前端收集参数发送POST /api/v1/robot/laser/command请求，后端为每条指令生成唯一的指令ID（格式为cmd-YYYYMMDD-XXX），记录指令动作、参数JSON、状态（待执行）和创建时间，保存到laser_command表。')

    add_text('树莓派端的指令轮询线程每1-2秒调用GET /api/v1/robot/commands/pending接口查询待执行指令，获取指令后模拟执行过程，根据指令类型设置不同的执行时长：FIRE指令根据照射时长参数sleep相应时间，SELF_TEST指令sleep 3秒模拟自检过程，其他指令sleep 500毫秒快速完成。执行完成后生成执行结果和消息，调用POST /api/v1/robot/laser/feedback上报执行反馈，调用PUT /api/v1/robot/commands/{id}/ack确认指令已执行。后端更新指令状态为已完成，创建激光操作日志记录，并通过WebSocket广播执行反馈信息，前端收到后更新激光控制页面的日志列表和设备状态。')

    add_text('激光日志页面提供激光设备操作日志的查询和筛选功能，布局与状态历史页面类似，用户可以按时间范围、操作类型、执行结果等条件筛选日志，查询结果支持分页显示、导出CSV和批量删除操作。')

    add_heading3('5.2.5 杂草检测模块')

    add_text('杂草检测模块负责检测结果的上报、存储、查询以及在线检测功能，整合了YOLOv8姿态估计模型推理服务。杂草检测页面提供在线图片检测功能，采用三栏布局：')

    add_text('（1）左侧图片上传区域：支持拖拽上传和点击选择文件，拖拽区域有虚线边框和上传提示文字，支持JPG、PNG格式图片。上传后显示图片缩略图和文件名、大小信息。')

    add_text('（2）中部参数设置区域：提供置信度滑块（范围0.01-1.0，默认0.25）和类别选择下拉框（全部、作物、杂草），用户可以根据需要调整检测参数，参数改变后自动重新推理显示结果。')

    add_text('（3）右侧结果展示区域：并排显示原图和标注图两张图片，支持鼠标滚轮缩放查看，标注图上用不同颜色的边界框和关键点标识检测到的作物和杂草。作物用绿色框标注，杂草用红色框标注，边界框左上角标注类别名称和置信度百分比，杂草茎部关键点用红色圆点标记。')

    add_text('（4）底部检测统计区域：显示总目标数、杂草数量、作物数量、推理耗时（毫秒）等信息，用大号字体突出显示。')

    add_text('（5）底部目标列表区域：使用Element Plus表格展示每个检测目标的详细信息，包括序号、类别、置信度、像素坐标、关键点信息、深度值、三维坐标等列。表格支持按类别筛选，方便用户查看特定类别的检测目标。')

    add_text('YOLO在线检测功能实现了图片上传、模型推理和结果解析。用户上传图片后，前端通过FormData格式发送POST /api/v1/detection/infer请求，后端生成唯一的请求ID，保存图片到指定目录，然后通过ProcessBuilder执行Python推理脚本调用Ultralytics YOLOv8姿态估计模型。推理完成后，Python脚本输出JSON格式的推理结果，包含每个检测目标的类别ID、类别名称、置信度、边界框坐标、关键点坐标等信息。后端解析JSON结果，使用OpenCV在原图上绘制检测框和关键点生成标注图片，返回推理结果JSON、标注图片URL、原始图片URL和运行日志给前端。')

    add_text('机载仿真端的检测上报线程每5-10秒从测试图片库中随机选择一张图片，调用后端YOLO检测接口进行推理，获取检测结果后，将标注图片和检测目标列表通过POST /api/v1/detection/report接口上报给后端系统。上报的检测结果包含原始图片、标注图片、杂草数量、作物数量、推理耗时和每个检测目标的详细信息。后端保存图片到本地文件系统，路径格式为files/detection/raw/日期/文件名和files/detection/result/日期/文件名，保存检测记录到detection_record表，保存检测目标到detection_item表。')

    add_text('检测记录页面提供杂草检测记录的查询和筛选功能，布局与状态历史页面类似。用户可以按时间范围、检测类别等条件筛选记录，查询结果支持分页显示，每条记录显示缩略图、检测时间、杂草数量、作物数量、推理耗时等信息。点击记录可以跳转到检测详情页面，查看完整的标注图片、检测统计、目标列表等信息，详情页面布局与YOLO在线检测页面类似。')

    add_heading3('5.2.6 WebSocket实时推送模块')

    add_text('WebSocket实时推送模块负责将机器人状态更新、检测结果、指令执行反馈等信息实时推送给前端。后端使用Spring WebSocket实现WebSocket服务，自定义WebSocketHandler处理WebSocket连接和消息。WebSocketHandler维护一个在线客户端会话列表，当有新的状态更新、检测结果或指令执行反馈时，遍历所有在线会话，通过WebSocket广播消息给所有连接的客户端。广播的消息类型包括STATUS_UPDATE（状态更新）、DETECTION_RESULT（检测结果）、COMMAND_FEEDBACK（指令执行反馈）等，消息体包含相应的数据内容。前端收到消息后根据消息类型更新相应的页面数据。')

    add_text('前端WebSocket通信工具封装在src/utils/websocket.js中，管理WebSocket连接的建立、消息监听、重连逻辑和关闭。前端在登录成功后建立WebSocket连接，连接地址为ws://{host}/ws/robot-status。当WebSocket连接断开时，系统会自动尝试重连，最多重连10次，每次重连间隔3秒，采用指数退避策略延长重连间隔。前端通过监听不同类型的消息，实时更新仪表盘的状态卡片和图表、激光控制页面的执行日志等页面数据，保证用户看到的信息始终是最新的。')

    add_heading3('5.2.7 机载仿真端模块')

    add_text('机载仿真端负责模拟真实机器人的运行行为，包括状态上报、杂草检测、指令执行等功能，采用多线程架构，使用Python语言实现。机载仿真端使用requests库进行HTTP通信，使用threading模块实现多线程，使用queue模块实现线程安全队列，使用rich库实现终端实时可视化面板，展示仿真端的运行状态和统计数据。')

    add_text('机载仿真端的多线程架构包括主线程、状态上报线程、检测上报线程、指令轮询线程：')

    add_text('（1）主线程：负责加载config.json配置文件（包含服务端地址、用户名密码、上报间隔参数等），调用登录接口获取认证Token，初始化HTTP会话headers，创建三个工作线程并启动，注册SIGINT信号处理函数处理Ctrl+C退出信号，使用rich库的Live面板实时显示各线程的运行状态和统计数据，主线程循环等待退出信号，收到后设置停止事件，等待所有工作线程安全退出后释放资源并退出程序。')

    add_text('（2）状态上报线程：每5秒执行一次上报任务。线程维护状态变量的当前值，每次上报时在合理范围内随机变化这些参数：电量在0-100%范围内缓慢下降（每次减少0.05-0.2%），低于20%时模拟充电回升；温度在35-50℃范围内波动；CPU使用率在20-80%范围内波动；GPS坐标每次增加小的随机偏移模拟机器人前进；速度在0.3-1.2 m/s范围内波动；IMU数据加入小的高斯噪声模拟传感器真实输出。状态数据组装成JSON后调用POST /api/v1/robot/status接口上报，更新上报成功/失败统计计数。')

    add_text('（3）检测上报线程：每5-10秒（随机间隔）执行一次检测任务。线程从test_images目录随机选择一张JPG图片，调用POST /api/v1/detection/infer接口上传图片执行YOLO检测推理，获取推理结果JSON和标注图片URL。然后组装multipart/form-data格式数据，包含原始图片文件、标注图片文件和检测结果JSON，调用POST /api/v1/detection/report接口上报检测结果，更新检测次数和成功次数统计。')

    add_text('（4）指令轮询线程：每1-2秒（随机间隔）调用GET /api/v1/robot/commands/pending接口查询待执行指令。获取到指令列表后遍历每条指令，根据指令类型模拟执行过程：FIRE指令sleep duration参数指定的毫秒数；SELF_TEST指令sleep 3秒模拟自检；ENABLE/DISABLE/STOP/RESET指令sleep 0.5秒；SET_POWER指令sleep 1秒；AIM指令sleep 0.8秒。执行完成后生成执行结果（90%概率SUCCESS，5%概率FAILED，5%概率TIMEOUT）和相应的执行消息，调用POST /api/v1/robot/laser/feedback上报执行结果，调用PUT /api/v1/robot/commands/{id}/ack确认指令已执行，更新指令执行统计。')

    add_heading2('5.3 系统测试')

    add_heading3('5.3.1 测试环境')

    add_text('系统测试环境配置如下：操作系统为Windows 11 22H2或macOS Ventura 13.5+；浏览器为Chrome 120+、Edge 120+或Firefox 115+；Java版本为OpenJDK 17.0.8；Node.js版本为18.18.0 LTS；MySQL版本为8.0.35；Python版本为3.10.12。测试硬件配置为Intel Core i7-12700H处理器，16GB DDR4内存，512GB NVMe固态硬盘。')

    add_heading3('5.3.2 功能测试')

    add_text('功能测试对系统的各项功能进行验证，主要测试用例及结果如下：')

    add_text('（1）用户认证模块：测试用户登录功能，输入正确的用户名admin和密码123456，验证成功登录并跳转到首页；测试错误密码登录，输入错误密码验证提示"用户名或密码错误"且不跳转；测试Token过期验证，修改系统时间24小时后验证访问接口返回401并自动跳转到登录页；测试用户登出功能，点击退出登录验证成功跳转到登录页且令牌失效。测试结果表明，用户认证模块功能正常，安全性良好。')

    add_text('（2）机器人状态监测模块：测试仪表盘实时状态显示，启动机载仿真端后验证状态卡片数据每5秒更新一次；测试图表更新，验证电量、CPU、温度折线图随状态上报正确更新数据点；测试统计数据显示，验证今日检测次数等统计数据随检测上报正确累加；测试状态历史查询，选择不同时间范围验证返回相应时间段的记录；测试筛选功能，设置电量<20%验证正确筛选出低电量记录；测试排序功能，点击温度列表头验证按温度升序/降序排列；测试导出功能，点击导出按钮验证成功下载CSV文件且内容正确。测试结果表明，仪表盘能够实时更新状态数据，图表数据正确更新，统计数据准确显示，状态历史模块的筛选功能完善，排序正确，导出功能正常。')

    add_text('（3）轨迹管理模块：测试轨迹显示功能，选择包含GPS数据的时间段验证地图上正确绘制轨迹线；测试轨迹点点击，点击轨迹上的点验证弹出正确的信息窗口；测试轨迹回放，点击播放按钮验证机器人标记沿轨迹移动且信息同步更新；测试速度调节，切换不同倍速验证回放速度相应变化。测试结果表明，轨迹地图能够正确绘制机器人运动轨迹，轨迹回放功能正常。')

    add_text('（4）激光设备控制模块：测试指令下发功能，点击"开启激光"按钮验证提示"指令已下发"且出现在待执行列表；测试指令执行，启动机载仿真端等待1-2秒验证指令状态变为已完成且出现在执行日志中；测试状态查询，刷新页面验证激光状态与指令执行结果一致；测试日志查询，切换到激光日志页面验证显示所有历史操作记录。测试结果表明，激光控制模块能够正确下发指令，查询设备状态，记录操作日志。')

    add_text('（5）杂草检测模块：测试图片上传功能，选择本地图片验证成功上传并显示缩略图；测试模型推理，点击"开始检测"按钮验证几秒后显示标注图片和检测结果；测试参数调整，拖动置信度滑块验证检测结果随之变化；测试类别过滤，选择"仅杂草"验证只显示杂草类别的检测目标；测试检测记录查询，切换到检测记录页面验证显示所有上报的检测记录；测试检测详情，点击一条记录验证跳转到详情页且显示完整的检测信息。测试结果表明，杂草检测模块能够正确上传图片，调用模型推理，参数调整后检测结果随之变化，检测记录查询功能正常。')

    add_text('（6）WebSocket模块：测试实时推送功能，启动机载仿真端后不刷新页面验证仪表盘状态自动更新；测试检测结果推送，机载仿真端上报检测后验证仪表盘统计数字自动增加；测试指令执行反馈推送，下发激光指令后验证前端自动收到执行结果并更新日志。测试结果表明，WebSocket能够实时推送状态更新、检测结果和指令执行反馈。')

    add_text('（7）机载仿真端模块：测试状态上报功能，启动仿真端验证后端每秒收到状态上报请求且数据库中新增状态记录；测试检测上报功能，启动仿真端验证后端收到检测上报请求且数据库中新增检测记录；测试指令执行功能，前端下发指令验证仿真端收到指令并上报执行反馈。测试结果表明，机载仿真端能够正常模拟机器人运行行为，上报状态数据和检测结果，执行指令并上报反馈。')

    add_heading3('5.3.3 性能测试')

    add_text('性能测试主要测试系统的响应时间和并发处理能力。使用JMeter 5.6工具进行测试，模拟不同并发用户数访问系统接口，测量平均响应时间、95%响应时间和成功率。测试结果如下：')

    add_text('（1）用户登录接口：在50并发用户下，平均响应时间为120ms，95%响应时间为180ms，成功率为100%。密码BCrypt验证是主要耗时点，但在可接受范围内。')

    add_text('（2）查询最新状态接口：在100并发用户下，平均响应时间为45ms，95%响应时间为80ms，成功率为100%。该接口只需单条数据库查询，性能优秀。')

    add_text('（3）查询状态历史接口（返回100条记录）：在50并发用户下，平均响应时间为150ms，95%响应时间为250ms，成功率为100%。分页查询性能良好。')

    add_text('（4）上报机器人状态接口：在100并发用户下，平均响应时间为60ms，95%响应时间为120ms，成功率为100%。单条数据插入操作性能优秀。')

    add_text('（5）上报检测结果接口（包含2MB图片上传）：在30并发用户下，平均响应时间为350ms，95%响应时间为500ms，成功率为100%。文件IO是主要耗时点，性能符合预期。')

    add_text('（6）YOLO在线检测接口（单张图片推理）：在10并发用户下，平均响应时间为1200ms，95%响应时间为1800ms，成功率为100%。YOLO模型推理是主要耗时点，性能符合AI推理服务的预期范围。')

    add_text('测试结果表明，系统在100并发用户下的响应时间均在可接受范围内，各项接口的成功率达到100%，系统具有良好的性能表现，满足设计的性能需求。系统整体运行稳定，没有出现内存泄漏或服务崩溃情况。')

    add_heading2('5.4 本章小结')

    add_text('本章详细介绍了激光除草机器人运行监测系统的实现与测试。首先介绍了开发环境的搭建，包括后端和前端的技术栈、开发工具、依赖版本和配置文件说明。然后详细阐述了各功能模块的实现细节和界面设计，包括用户认证与授权模块的登录登出流程和界面、机器人状态监测模块的仪表盘和状态历史页面、轨迹管理模块的地图展示和回放功能、激光设备控制模块的指令下发和执行流程、杂草检测模块的在线检测和记录管理、WebSocket实时推送模块的前后端实现、机载仿真端的多线程架构和实现。最后进行了系统测试，包括功能测试、性能测试，详细列出了各项测试用例和测试结果。测试结果表明系统功能完善、性能良好，达到了预期的设计目标。')

    doc.add_page_break()

    # ==================== 总结与展望 ====================
    add_heading1('总结与展望')

    add_heading2('工作总结')

    add_text('本课题设计并实现了一套基于Web的激光除草机器人运行监测系统，采用三层B/S架构，通过软件仿真的方式在不依赖实际硬件设备的情况下，完整验证了激光除草机器人监控系统的业务流程与技术可行性。系统实现了机器人状态实时监控、检测记录管理、激光控制指令管理、轨迹地图可视化、YOLO在线检测等核心功能，界面简洁直观，操作便捷，能够有效帮助操作人员远程监控机器人的作业状态。')

    add_text('本课题的主要工作成果包括：')

    add_text('（1）完成了系统的需求分析，详细分析了激光除草机器人监控系统的功能需求与非功能需求，给出了系统用例图。')

    add_text('（2）完成了系统的架构设计，采用三层B/S架构设计方案，设计了系统的核心工作流程、数据库表结构、核心模块类图、机载仿真端多线程架构和通信机制，给出了相应的架构图。')

    add_text('（3）完成了后端服务的开发，基于Spring Boot框架实现了用户认证、机器人状态管理、检测记录管理、激光控制管理、统计分析、WebSocket实时推送等核心业务模块。')

    add_text('（4）完成了前端系统的开发，基于Vue 3框架和Element Plus组件库实现了登录页面、仪表盘页面、状态历史页面、检测记录页面、激光控制页面、轨迹地图页面、YOLO在线检测页面等功能页面，使用ECharts实现数据可视化。')

    add_text('（5）完成了YOLO模型的集成调用，集成了YOLOv8姿态估计模型，实现了图片上传、模型推理、结果解析、标注图片生成的完整在线检测流程。本系统仅负责模型集成调用，不涉及模型训练和硬件加速工作。')

    add_text('（6）完成了机载仿真端的开发，基于Python和多线程技术实现了机器人状态上报、杂草检测上报、激光指令轮询执行等模拟功能，通过软件仿真验证了系统功能。')

    add_text('（7）完成了系统测试，对系统进行了全面的功能测试和性能测试，测试结果表明系统功能完善、性能良好、运行稳定，满足设计需求。')

    add_heading2('存在的不足')

    add_text('虽然本系统基本实现了预期的功能目标，但仍存在一些不足之处，主要包括：')

    add_text('（1）未接入真实硬件设备：当前系统通过软件仿真的方式模拟机器人行为，尚未在真实的激光除草机器人硬件平台上进行部署和测试，实际应用效果有待验证。')

    add_text('（2）YOLO模型未进行针对性优化：当前系统直接使用通用YOLOv8模型，未针对杂草检测场景进行专门的数据集训练和模型优化，检测精度和速度还有提升空间。')

    add_text('（3）系统功能还可扩展：当前系统实现了基本的监控功能，但缺少告警通知、报表统计、多机器人管理、远程调试等高级功能，功能完整性还有待增强。')

    add_text('（4）前端移动端适配不足：当前前端页面主要面向桌面端设计，在手机等移动设备上的显示效果和操作体验还有待优化。')

    add_heading2('未来展望')

    add_text('针对上述不足之处，未来可以从以下几个方面进行改进和扩展：')

    add_text('（1）接入真实硬件平台：将系统部署到真实的激光除草机器人硬件平台上，对接真实的树莓派、深度相机、激光发射器等硬件，实现对真实除草机器人的监控管理，在实际田间作业中验证系统的实用价值。')

    add_text('（2）优化杂草检测模型：收集田间真实的杂草和作物图像，构建专用的杂草检测数据集，针对农业场景对YOLO模型进行训练优化，提升复杂田间环境下的杂草检测精度和速度。同时研究模型轻量化技术，使其能够更好地在边缘设备上运行。')

    add_text('（3）扩展系统高级功能：增加异常告警功能，当机器人电量过低、温度过高、检测到故障时通过短信、邮件等方式及时通知操作人员；增加报表统计功能，支持按日、周、月生成作业统计报表，支持PDF导出；增加多机器人管理功能，支持同时监控多台除草机器人；增加远程调试功能，支持远程查看机载端日志、远程配置参数等。')

    add_text('（4）优化前端移动端适配：采用响应式设计优化前端页面，适配手机、平板等移动设备，开发移动端专用App或小程序，方便用户随时随地监控机器人状态。')

    add_text('（5）增加数据智能分析：利用机器学习算法对历史检测数据进行分析，识别杂草分布规律，为田间管理提供决策支持；分析机器人运行状态数据，进行故障预测和健康管理。')

    add_text('（6）提升系统安全等级：增加操作审计日志，记录所有用户的操作行为；增加数据加密存储，保护敏感数据；增加防SQL注入、防XSS攻击等安全加固措施，提升系统整体安全性。')

    doc.add_page_break()

    # ==================== 参考文献 ====================
    add_heading1('参考文献')

    add_ref('[1] Oerke E C. Crop losses to pests[J]. The Journal of Agricultural Science, 2006, 144(1): 31-43.')

    add_ref('[2] Gianessi L P. The increasing importance of herbicides in worldwide crop production[J]. Pest Management Science, 2013, 69(10): 1099-1105.')

    add_ref('[3] Hamuda E, Glavin M, Jones E. A survey of image processing techniques for plant extraction and segmentation in the field[J]. Computers and Electronics in Agriculture, 2016, 125: 184-199.')

    add_ref('[4] Heisel T, Andreasen C, Ascard J. Weed control by light radiation—A review[J]. Weed Research, 2001, 41(5): 357-372.')

    add_ref('[5] Craessaerts G, Saeys W, Van Eetvelde V, et al. An innovative instrument for automatic weed detection and selective laser weeding[J]. Proceedings of the 6th European Conference on Precision Agriculture, 2007: 305-312.')

    add_ref('[6] 王金龙, 王建国. 基于Web的远程监控系统设计与实现[J]. 计算机工程与设计, 2011, 32(6): 1985-1988.')

    add_ref('[7] Rakhmatulin I, Andreasen C. A review of weed detection and identification technologies[J]. AgriEngineering, 2020, 2(2): 277-305.')

    add_ref('[8] Carbon Robotics. Autonomous LaserWeeder[EB/OL]. https://carbonrobotics.com/, 2023.')

    add_ref('[9] Blue River Technology. See & Spray Technology[EB/OL]. https://bluerivertechnology.com/, 2023.')

    add_ref('[10] 张伟, 李民赞, 张俊逸, 等. 基于改进YOLOv5的田间杂草实时检测方法[J]. 农业机械学报, 2022, 53(4): 245-253.')

    add_ref('[11] 陈树人, 张宝峰, 许晓东, 等. 基于机器视觉的激光除草系统设计与试验[J]. 农业工程学报, 2020, 36(12): 157-164.')

    add_ref('[12] John Deere. John Deere Operations Center[EB/OL]. https://www.deere.com/en/technology-and-innovation/operations-center/, 2023.')

    add_ref('[13] 王辉, 吕钊钦, 李法德, 等. 智能农机远程监控系统研究进展与趋势[J]. 农业机械学报, 2021, 52(10): 1-15.')

    add_ref('[14] You Y. Vue.js: A progressive JavaScript framework[EB/OL]. https://vuejs.org/, 2023.')

    add_ref('[15] Pivotal Software. Spring Boot[EB/OL]. https://spring.io/projects/spring-boot, 2023.')

    add_ref('[16] Fette I, Melnikov A. The WebSocket Protocol[R]. RFC 6455, 2011.')

    add_ref('[17] You Y. Vue 3 Documentation[EB/OL]. https://vuejs.org/guide/introduction.html, 2023.')

    add_ref('[18] You Y. Vite - Next Generation Frontend Tooling[EB/OL]. https://vitejs.dev/, 2023.')

    add_ref('[19] Element Plus Team. Element Plus[EB/OL]. https://element-plus.org/, 2023.')

    add_ref('[20] Apache Software Foundation. Apache ECharts[EB/OL]. https://echarts.apache.org/, 2023.')

    add_ref('[21] Pivotal Software. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/current/reference/html/, 2023.')

    add_ref('[22] Baomidou. MyBatis-Plus[EB/OL]. https://baomidou.com/, 2023.')

    add_ref('[23] Jones M, Bradley J, Sakimura N. JSON Web Token (JWT)[R]. RFC 7519, 2015.')

    add_ref('[24] Oracle Corporation. MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/, 2023.')

    add_ref('[25] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.')

    add_ref('[26] Redmon J, Farhadi A. YOLOv3: An Incremental Improvement[EB/OL]. https://pjreddie.com/media/files/papers/YOLOv3.pdf, 2018.')

    add_ref('[27] Ultralytics. YOLOv8 Documentation[EB/OL]. https://docs.ultralytics.com/, 2023.')

    doc.add_page_break()

    # ==================== 致谢 ====================
    add_heading1('致谢')

    add_text('时光荏苒，四年的大学生活即将结束。回首这段难忘的求学时光，心中充满了感激之情。在论文完成之际，我谨向所有给予我帮助和支持的人们表示最诚挚的感谢。')

    add_text('首先，我要感谢我的指导教师XXX老师。在毕业设计的整个过程中，X老师给予了我悉心的指导和无私的帮助。从选题、开题、系统设计到代码实现、论文撰写的各个阶段，X老师都提出了宝贵的意见和建议，使我能够顺利完成本次毕业设计。X老师严谨的治学态度、渊博的专业知识、诲人不倦的师者风范和敬业的工作精神，都给我留下了深刻的印象，也将对我今后的学习和工作产生深远的影响。再次向X老师表示最衷心的感谢！')

    add_text('感谢实验室的同学们，在毕业设计的过程中，我们相互交流、相互帮助，共同克服了许多困难。特别感谢XXX同学在前端Vue开发方面给予的帮助，感谢XXX同学在后端Spring Boot开发中分享的经验，感谢XXX同学在YOLO模型集成方面给予的指导。与同学们的交流讨论让我受益匪浅，也使我感受到了团队合作的重要性。')

    add_text('感谢我的室友们，四年来我们朝夕相处，相互关心，共同进步。感谢你们在生活上的照顾和学习上的帮助，和你们在一起的日子是我大学四年最美好的回忆。')

    add_text('感谢我的家人和朋友们，他们一直以来都给予我最无私的关心和支持，在我遇到困难的时候鼓励我，在我取得进步的时候为我高兴。你们的支持是我不断前进的动力，没有你们就没有今天的我。')

    add_text('最后，感谢所有在毕业设计过程中给予我帮助的人们，感谢评阅本论文的各位专家和老师，谢谢你们！')

    # 保存文档
    doc.save('激光除草运行监测系统毕业论文_修改版.docx')
    print('论文修改完成，已保存为：激光除草运行监测系统毕业论文_修改版.docx')

if __name__ == '__main__':
    create_revised_thesis()
