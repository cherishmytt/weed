from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========== 1. 任务书 ==========
doc1 = Document()

# 设置标题
title = doc1.add_heading('福州大学本科生毕业设计（论文）任务书', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加基本信息表
doc1.add_paragraph('题目：基于YOLOv8-Pose的农田杂草根茎定位检测系统研究与实现')
doc1.add_paragraph('')
doc1.add_paragraph('学号：102201240')
doc1.add_paragraph('姓名：')
doc1.add_paragraph('专业：网络工程与技术')
doc1.add_paragraph('学院：计算机与大数据学院')
doc1.add_paragraph('')

# 任务内容
doc1.add_heading('一、主要内容及基本要求', level=1)

doc1.add_paragraph('（一）主要研究内容：')
doc1.add_paragraph('1. 研究农田杂草检测技术现状，分析基于深度学习的目标检测算法在农业领域的应用前景。')
doc1.add_paragraph('2. 研究YOLOv8-Pose姿态估计算法原理，探索其在杂草根茎关键点定位中的应用。')
doc1.add_paragraph('3. 构建CropAndWeed数据集到YOLO-Pose格式的转换工具，实现杂草根茎关键点标注的自动转换。')
doc1.add_paragraph('4. 基于YOLOv8s-Pose模型训练杂草检测模型，实现作物与杂草的分类及根茎精准定位。')
doc1.add_paragraph('5. 开发基于Spring Boot的Web可视化系统，实现图片上传、模型推理、检测结果可视化展示功能。')
doc1.add_paragraph('6. 设计并实现REST API接口，支持第三方系统调用杂草检测服务。')

doc1.add_paragraph('')
doc1.add_paragraph('（二）基本要求：')
doc1.add_paragraph('1. 完成相关文献阅读，撰写不少于5000字的文献综述。')
doc1.add_paragraph('2. 完成一篇外文文献翻译，不少于20000外文印刷字符。')
doc1.add_paragraph('3. 实现数据转换工具，支持CropAndWeed数据集到YOLO-Pose格式转换。')
doc1.add_paragraph('4. 训练得到mAP@0.5不低于0.85的杂草检测模型。')
doc1.add_paragraph('5. Web系统支持图片上传、参数配置、结果可视化、坐标导出功能。')
doc1.add_paragraph('6. 撰写不少于15000字的毕业论文。')

doc1.add_heading('二、研究进度安排', level=1)
progress = [
    ('第1-2周', '查阅文献，确定选题，完成开题报告'),
    ('第3-4周', '环境搭建，数据集获取与预处理，完成数据转换工具'),
    ('第5-8周', 'YOLOv8-Pose模型训练，参数调优，完成模型验证'),
    ('第9-12周', 'Spring Boot Web系统开发，前后端联调，API接口实现'),
    ('第13-14周', '系统集成测试，性能优化，功能完善'),
    ('第15-16周', '撰写毕业论文，修改完善，准备答辩'),
]
for week, content in progress:
    p = doc1.add_paragraph()
    p.add_run(f'{week}：').bold = True
    p.add_run(content)

doc1.add_heading('三、主要参考文献', level=1)
refs = [
    '[1] Redmon J, Farhadi A. YOLOv3: An Incremental Improvement[EB/OL]. arXiv:1804.02767, 2018.',
    '[2] Ultralytics. YOLOv8 Documentation[EB/OL]. https://docs.ultralytics.com/, 2023.',
    '[3] Haug S, Ostermann J. A Crop/Weed Field Image Dataset for the Evaluation of Computer Vision Based Precision Agriculture Tasks[C]. CVPPP 2014.',
    '[4] You J, Liu X, Li Y. Real-time weed detection in outdoor field conditions using deep learning[J]. Computers and Electronics in Agriculture, 2021, 186: 106212.',
    '[5] Wang A, Zhang W, Wei X. A review on weed detection using ground-based machine vision and image processing techniques[J]. Computers and Electronics in Agriculture, 2019, 158: 226-240.',
]
for ref in refs:
    doc1.add_paragraph(ref, style='List Number')

# 保存
doc1.save('1.福州大学本科生毕业设计（论文）任务书-已完成.docx')
print('任务书生成完成')

# ========== 2. 开题报告 ==========
doc2 = Document()
title2 = doc2.add_heading('福州大学本科生毕业设计（论文）开题报告', 0)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc2.add_paragraph('题目：基于YOLOv8-Pose的农田杂草根茎定位检测系统研究与实现')
doc2.add_paragraph('')
doc2.add_paragraph('学号：102201240')
doc2.add_paragraph('姓名：')
doc2.add_paragraph('专业：网络工程与技术')
doc2.add_paragraph('学院：计算机与大数据学院')
doc2.add_paragraph('')

doc2.add_heading('一、研究背景与意义', level=1)
doc2.add_paragraph('随着全球农业可持续发展需求的日益增长，精准农业技术成为提高农业生产效率、减少化学农药使用的关键方向。杂草作为影响作物产量的重要因素，传统的化学除草和机械除草方法存在环境污染、成本高昂、作物损伤等问题。')
doc2.add_paragraph('近年来，基于计算机视觉和深度学习的智能除草技术迅速发展，其中激光除草以其精准、高效、无污染的特点成为研究热点。精准定位杂草根茎是实现激光除草的核心技术需求，只有准确识别杂草并定位其根茎位置，才能实现精准打击，避免误伤作物。')
doc2.add_paragraph('本研究利用YOLOv8-Pose姿态估计算法，通过检测杂草根茎关键点，实现作物与杂草的分类识别和根茎精准定位，为智能激光除草设备提供视觉感知能力。该研究对推动农业智能化发展、减少化学除草剂使用、实现绿色农业具有重要的理论意义和应用价值。')

doc2.add_heading('二、国内外研究现状', level=1)
doc2.add_paragraph('在目标检测领域，YOLO系列算法凭借其速度和精度的平衡成为实时检测的首选方案。YOLOv8在检测精度和速度上都有显著提升，并支持姿态估计任务。姿态估计原本主要用于人体关键点检测，但其"边界框+关键点"的输出形式与杂草检测中"目标识别+根茎定位"的需求天然契合。')
doc2.add_paragraph('在农业杂草检测方面，已有大量研究利用深度学习实现杂草识别。CropAndWeed数据集提供了包含作物和杂草边界框及根茎点的标注数据，为基于关键点的杂草检测提供了数据基础。现有研究多集中于杂草的分类识别，对于根茎精准定位的研究相对较少。')
doc2.add_paragraph('在智能除草系统方面，Carbon Robotics、Blue River Technology等公司已推出商业激光除草设备，但设备价格昂贵，且核心算法不开源。针对农田复杂场景下的小目标检测、遮挡处理、实时性要求等仍是亟待解决的问题。')

doc2.add_heading('三、主要研究内容', level=1)
contents = [
    '1. 数据集预处理与格式转换：研究CropAndWeed数据集结构，设计并实现数据转换脚本，将原始CSV标注转换为YOLO-Pose格式的关键点标注，完成数据集划分。',
    '2. YOLOv8-Pose模型训练与优化：基于YOLOv8s-Pose预训练模型，在CropAndWeed数据集上进行迁移学习，调整训练参数（batch size、learning rate、epochs等），优化模型性能。',
    '3. 模型性能评估：通过mAP、精确度、召回率、OKS（目标关键点相似度）等指标评价模型性能，分析不同参数对检测结果的影响。',
    '4. Web可视化系统开发：基于Spring Boot框架开发Web系统，实现图片上传、模型推理调用、检测结果可视化（边界框、关键点标注）、坐标明细展示、API接口等功能。',
    '5. 系统集成与测试：完成模型与Web系统的集成，进行功能测试和性能优化，确保系统稳定运行。'
]
for c in contents:
    doc2.add_paragraph(c)

doc2.add_heading('四、研究方案与技术路线', level=1)
doc2.add_paragraph('本研究采用理论研究与工程实现相结合的方法，具体技术路线如下：')
doc2.add_paragraph('1. 理论学习阶段：深入学习YOLO系列算法原理、姿态估计方法、Spring Boot框架等相关技术。')
doc2.add_paragraph('2. 数据准备阶段：获取CropAndWeed数据集，编写Python脚本实现数据格式转换，构建YOLO-Pose训练数据集。')
doc2.add_paragraph('3. 模型训练阶段：使用Ultralytics库加载YOLOv8s-pose预训练模型，在准备好的数据集上进行训练，通过调整训练参数优化模型性能。')
doc2.add_paragraph('4. 系统开发阶段：设计Web系统架构，使用Spring Boot + Vue.js开发前后端，实现文件上传、Python脚本调用、结果渲染等功能。')
doc2.add_paragraph('5. 测试与完善阶段：进行系统功能测试、性能测试，修复存在的问题，优化用户体验。')

doc2.add_heading('五、预期目标和成果', level=1)
doc2.add_paragraph('1. 构建适用于YOLO-Pose训练的杂草根茎关键点检测数据集。')
doc2.add_paragraph('2. 训练得到一个能够同时检测作物/杂草类别和根茎关键点的YOLOv8-Pose模型，mAP@0.5达到0.85以上。')
doc2.add_paragraph('3. 开发一个Web可视化检测系统，支持图片上传、参数配置、检测结果可视化展示。')
doc2.add_paragraph('4. 完成一篇符合学校要求的毕业论文。')

doc2.add_heading('六、进度安排', level=1)
for week, content in progress:
    p = doc2.add_paragraph()
    p.add_run(f'{week}：').bold = True
    p.add_run(content)

doc2.add_heading('七、参考文献', level=1)
for ref in refs:
    doc2.add_paragraph(ref, style='List Number')

doc2.save('2.福州大学本科生毕业设计（论文）开题报告-已完成.docx')
print('开题报告生成完成')

# ========== 3. 中期检查表 ==========
doc3 = Document()
title3 = doc3.add_heading('福州大学本科生毕业设计（论文）中期检查表', 0)
title3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc3.add_paragraph('题目：基于YOLOv8-Pose的农田杂草根茎定位检测系统研究与实现')
doc3.add_paragraph('')
doc3.add_paragraph('学号：102201240')
doc3.add_paragraph('姓名：')
doc3.add_paragraph('专业：网络工程与技术')
doc3.add_paragraph('学院：计算机与大数据学院')
doc3.add_paragraph('')

doc3.add_heading('一、任务完成情况', level=1)
doc3.add_paragraph('1. 文献综述：已完成相关文献阅读，阅读了YOLO系列算法、姿态估计、农业计算机视觉等领域的中英文文献20余篇，完成了文献综述撰写。')
doc3.add_paragraph('2. 外文翻译：完成一篇关于自动激光除草系统的外文文献翻译，字符数符合要求。')
doc3.add_paragraph('3. 环境搭建：已完成Python深度学习环境搭建，配置了CUDA 12.8、PyTorch、Ultralytics等依赖。')
doc3.add_paragraph('4. 数据准备：已获取CropAndWeed数据集，完成数据探索，编写了cnw2yolo.py转换脚本，成功将数据集转换为YOLO-Pose格式。')
doc3.add_paragraph('5. 模型训练：已完成YOLOv8s-Pose模型训练，训练100 epoch，在验证集上取得了较好的mAP和OKS指标。')
doc3.add_paragraph('6. Web系统开发：已完成Spring Boot项目搭建，实现了图片上传、Python推理脚本调用、JSON结果返回、前端结果展示等核心功能。')

doc3.add_heading('二、已解决的关键问题', level=1)
doc3.add_paragraph('1. 解决了RTX 5060 Ti显卡（Blackwell架构）的CUDA兼容性问题，使用CUDA 12.8版本的PyTorch成功启用GPU加速。')
doc3.add_paragraph('2. 完成了CropAndWeed数据集到YOLO-Pose格式的转换，实现了边界框和关键点坐标的归一化处理。')
doc3.add_paragraph('3. 实现了Java调用Python脚本进行模型推理的方案，解决了跨语言调用的参数传递和结果解析问题。')
doc3.add_paragraph('4. 解决了Spring Boot项目的Java版本兼容性问题，统一配置为Java 17。')

doc3.add_heading('三、待解决的问题', level=1)
doc3.add_paragraph('1. 模型对小目标杂草的检测精度有待提升，需要调整输入分辨率或使用数据增强策略。')
doc3.add_paragraph('2. Web系统的用户界面需要进一步优化，增加更多交互功能。')
doc3.add_paragraph('3. 需要完善系统异常处理机制，提升系统稳定性。')
doc3.add_paragraph('4. 检测结果的批量导出功能尚未实现。')

doc3.add_heading('四、下一阶段工作计划', level=1)
plans = [
    '第10-11周：优化模型性能，尝试提高输入分辨率（960x960），调整数据增强参数，提升小目标检测效果。',
    '第12周：完善Web系统前端界面，优化用户体验，增加批量检测、结果导出等功能。',
    '第13周：进行系统集成测试，编写测试用例，修复发现的bug，完善异常处理机制。',
    '第14周：进行系统性能优化，包括推理速度优化、内存占用优化，撰写论文的系统实现部分。',
    '第15周：完成毕业论文撰写，进行论文查重和修改完善。',
    '第16周：准备答辩材料，制作PPT，进行答辩演练。'
]
for plan in plans:
    doc3.add_paragraph(plan)

doc3.add_heading('五、是否按计划进行', level=1)
doc3.add_paragraph('□ 提前完成    ■ 按计划进行    □ 延期')

doc3.add_heading('六、指导教师意见', level=1)
doc3.add_paragraph('签字：___________    日期：____年____月____日')

doc3.save('7.福州大学本科生毕业设计（论文）中期检查表-已完成.docx')
print('中期检查表生成完成')

print('\n所有文档生成完成！')
